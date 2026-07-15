import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import os
import textwrap
import datetime
import ollama
from ai_modules import (
    resume_matching,
    skill_gap_analyser,
    candidate_ranking,
    hiring_recommendation,
    resume_chat,
    chatbot_query,
    email_generator,
    recruitment_analysis,
    talent_insight,
    job_discussion_analyzer,
    interview_question_generator,
    talent_management_summary,
)

# 1. Must call set_page_config as the first Streamlit command
st.set_page_config(
    page_title="AI Recruitment & Talent Copilot",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded"
)

# Safe rerun function to support different Streamlit versions
def safe_rerun():
    try:
        st.rerun()
    except AttributeError:
        st.experimental_rerun()

# ==========================================
# WORKSPACE PAGE RENDERING FUNCTIONS
# ==========================================

def render_landing_page(candidates_df, requirements_df):
    st.markdown("### 🏢 Enterprise Talent Command Center")
    st.markdown("Welcome to the **AI Recruitment & Talent Copilot**. This dashboard aggregates recruitment performance metrics, candidate analytics, automated screening logs, and onboarding progression.")
    
    # KPIs Row
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(
            f'<div class="audit-metric-card">'
            f'<div class="audit-number">{len(candidates_df)}</div>'
            f'<div class="audit-label">Total Applicants</div>'
            f'</div>',
            unsafe_allow_html=True
        )
    with col2:
        st.markdown(
            f'<div class="audit-metric-card">'
            f'<div class="audit-number" style="color:#10b981;">{len(requirements_df)}</div>'
            f'<div class="audit-label">Job Openings</div>'
            f'</div>',
            unsafe_allow_html=True
        )
    with col3:
        st.markdown(
            f'<div class="audit-metric-card">'
            f'<div class="audit-number" style="color:#3b82f6;">{len(st.session_state["interviews"])}</div>'
            f'<div class="audit-label">Interviews Scheduled</div>'
            f'</div>',
            unsafe_allow_html=True
        )
    with col4:
        st.markdown(
            f'<div class="audit-metric-card">'
            f'<div class="audit-number" style="color:#a78bfa;">{len(st.session_state["onboarding_db"])}</div>'
            f'<div class="audit-label">Onboarding Pipelines</div>'
            f'</div>',
            unsafe_allow_html=True
        )
        
    st.write("")
    
    # Call-to-action Shortcuts
    st.markdown("#### ⚡ Quick Actions Dashboard")
    
    col_c1, col_c2, col_c3 = st.columns(3)
    
    with col_c1:
        st.markdown(
            """
            <div style="background-color: #1e293b; padding: 1.2rem; border-radius: 10px; border: 1px solid #334155; height: 180px;">
                <h5 style="color: #60a5fa; margin: 0 0 0.5rem 0; font-family: Sora;">Resume Center</h5>
                <p style="color: #cbd5e1; font-size: 0.85rem; line-height: 1.4;">Upload a candidate's resume and parse skills, experience, and certifications automatically using our simulated parser.</p>
            </div>
            """, unsafe_allow_html=True
        )
        if st.button("Go to Resume Uploader", key="go_upload", use_container_width=True):
            st.session_state["nav_option"] = "Resume Upload Page"
            safe_rerun()
            
    with col_c2:
        st.markdown(
            """
            <div style="background-color: #1e293b; padding: 1.2rem; border-radius: 10px; border: 1px solid #334155; height: 180px;">
                <h5 style="color: #34d399; margin: 0 0 0.5rem 0; font-family: Sora;">Job Architect</h5>
                <p style="color: #cbd5e1; font-size: 0.85rem; line-height: 1.4;">Generate industry-compliant job descriptions instantly and insert them directly into the recruitment matching system.</p>
            </div>
            """, unsafe_allow_html=True
        )
        if st.button("Go to JD Generator", key="go_jd", use_container_width=True):
            st.session_state["nav_option"] = "AI Job Description Generator"
            safe_rerun()
            
    with col_c3:
        st.markdown(
            """
            <div style="background-color: #1e293b; padding: 1.2rem; border-radius: 10px; border: 1px solid #334155; height: 180px;">
                <h5 style="color: #a78bfa; margin: 0 0 0.5rem 0; font-family: Sora;">Assessment Suite</h5>
                <p style="color: #cbd5e1; font-size: 0.85rem; line-height: 1.4;">Compare shortlisted candidates, inspect AI matching criteria trails, and view active bias audits.</p>
            </div>
            """, unsafe_allow_html=True
        )
        if st.button("Go to Assessments", key="go_assess", use_container_width=True):
            st.session_state["nav_option"] = "Candidate Assessment Suite"
            safe_rerun()
            
    st.write("")
    st.markdown("---")
    st.markdown("#### 📢 Recent Platform Notifications")
    for n in st.session_state["notifications"][:3]:
        read_dot = "🔵" if not n["Read"] else "⚪"
        st.markdown(f"{read_dot} **{n['Message']}** ({n['Time']})")

def render_candidate_portal(requirements_df):
    st.markdown("### 📝 Candidate Application Portal")
    st.markdown("Submit your application details below to apply for positions directly. The AI copilot will evaluate your profile against current requisitions.")
    
    with st.form("candidate_app_form", clear_on_submit=True):
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            name = st.text_input("Full Name", placeholder="e.g. Harish Kumar")
            email = st.text_input("Email Address", placeholder="e.g. harish@example.com")
        with col_f2:
            roles = requirements_df["Role"].tolist()
            applied_role = st.selectbox("Target Position", roles)
            experience = st.slider("Years of Experience", 0, 15, 2)
            
        skills_input = st.text_area("Key Skills (comma-separated)", placeholder="e.g. Python, SQL, Git, AWS")
        resume_summary = st.text_area("Profile Summary", placeholder="Brief description of your professional achievements...")
        
        submitted = st.form_submit_button("Submit Application", type="primary")
        
        if submitted:
            if not name.strip() or not skills_input.strip():
                st.error("Please enter your name and at least one skill to proceed.")
            else:
                new_cand = {
                    "Name": name.strip(),
                    "Role Applied": applied_role,
                    "Skills": skills_input.strip(),
                    "Experience_Years": float(experience),
                    "Match_Score": 0,
                    "Status": "Applied"
                }
                st.session_state["candidates_db"].append(new_cand)
                st.session_state["notifications"].insert(0, {
                    "Message": f"New application received from {name.strip()} for {applied_role}",
                    "Time": "Just now",
                    "Read": False
                })
                st.success(f"Thank you, {name}! Your application for **{applied_role}** has been submitted successfully.")

def render_resume_upload(requirements_df):
    st.markdown("### 📤 AI Resume Upload & Fit Analyzer")
    st.markdown("Upload a candidate resume, set the target job position, and instantly get an AI-powered resume fit analysis.")

    col_up1, col_up2 = st.columns([3, 2])
    with col_up1:
        uploaded_file = st.file_uploader("Upload Resume (PDF, TXT or DOCX)", type=["pdf", "txt", "docx"])
    with col_up2:
        target_role = st.selectbox("Target Job Position", requirements_df["Role"].tolist(), key="upload_target_role")

    # ---- Extract resume text ----
    resume_text = ""
    parsed_name = ""
    parsed_skills = ""
    parsed_exp = 3.0

    if uploaded_file is not None:
        file_ext = uploaded_file.name.lower().split(".")[-1]
        raw_bytes = uploaded_file.read()

        if file_ext == "txt":
            resume_text = raw_bytes.decode("utf-8", errors="ignore")
        elif file_ext == "pdf":
            try:
                import io
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
                    resume_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                except ImportError:
                    try:
                        import PyPDF2
                        reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
                        resume_text = "\n".join(
                            page.extract_text() or "" for page in reader.pages
                        )
                    except ImportError:
                        resume_text = f"[PDF content of {uploaded_file.name} — install pypdf for full text extraction]"
            except Exception as e:
                resume_text = f"[Could not extract PDF text: {e}]"
        elif file_ext == "docx":
            try:
                import io
                import docx
                doc = docx.Document(io.BytesIO(raw_bytes))
                resume_text = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                resume_text = f"[DOCX content of {uploaded_file.name} — install python-docx for full extraction]"
            except Exception as e:
                resume_text = f"[Could not extract DOCX text: {e}]"

        # Fall back to name-based demo data if extraction is empty
        fn = uploaded_file.name.lower()
        if not resume_text.strip():
            if "arun" in fn:
                resume_text = "Name: Arun. Skills: Python, SQL, Git, Docker, Communication. Experience: 4 years. Applied for Software Engineer."
                parsed_name, parsed_skills, parsed_exp = "Arun", "Python, SQL, Git, Docker, Communication", 4.0
            elif "divya" in fn:
                resume_text = "Name: Divya. Skills: Python, React, JavaScript, HTML, CSS. Experience: 2 years. Applied for Software Engineer."
                parsed_name, parsed_skills, parsed_exp = "Divya", "Python, React, JavaScript, HTML, CSS", 2.0
            else:
                parsed_name = uploaded_file.name.split(".")[0].replace("_", " ").title()
                parsed_skills = "Python, SQL, Git, Communication"
                parsed_exp = 3.0
                resume_text = f"Name: {parsed_name}. Skills: {parsed_skills}. Experience: {parsed_exp} years."
        else:
            parsed_name = uploaded_file.name.split(".")[0].replace("_", " ").title()

        st.success(f"✅ Resume **{uploaded_file.name}** loaded successfully ({len(resume_text)} characters extracted).")

        # ---- Editable parsed fields ----
        st.markdown("#### 📝 Review Extracted Details")
        with st.form("parsed_data_form"):
            c_name   = st.text_input("Candidate Name", value=parsed_name)
            c_role   = st.selectbox(
                "Role Applied For",
                requirements_df["Role"].tolist(),
                index=requirements_df["Role"].tolist().index(target_role)
                      if target_role in requirements_df["Role"].tolist() else 0
            )
            c_skills = st.text_input("Skills (comma separated)", value=parsed_skills or "Python, SQL, Git")
            c_exp    = st.number_input("Years of Experience", min_value=0.0, max_value=30.0,
                                       value=float(parsed_exp), step=0.5)

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                analyze_btn = st.form_submit_button("🤖 Analyze Resume Fit", type="primary", use_container_width=True)
            with col_btn2:
                save_btn = st.form_submit_button("💾 Confirm & Add to Candidate Pool", use_container_width=True)

            # ---- Save candidate ----
            if save_btn:
                new_cand = {
                    "Name": c_name.strip(),
                    "Role Applied": c_role,
                    "Skills":       c_skills.strip(),
                    "Experience_Years": float(c_exp),
                    "Match_Score":  0,
                    "Status":       "Screening"
                }
                st.session_state["candidates_db"] = [
                    c for c in st.session_state["candidates_db"] if c["Name"] != c_name.strip()
                ]
                st.session_state["candidates_db"].append(new_cand)
                st.session_state["notifications"].insert(0, {
                    "Message": f"Resume for {c_name} added to candidate pool.",
                    "Time": "Just now",
                    "Read": False
                })
                st.success(f"Candidate **{c_name}** added to pool!")
                st.session_state["nav_option"] = "Resume Analysis Report"
                st.session_state["last_analyzed_candidate"] = c_name.strip()
                safe_rerun()

            # ---- AI Resume Fit Analysis ----
            if analyze_btn:
                role_req = requirements_df[requirements_df["Role"] == c_role]
                if role_req.empty:
                    st.error("Could not find job requirements for the selected role.")
                else:
                    req_row      = role_req.iloc[0]
                    req_skills   = [s.strip().lower() for s in req_row["Required_Skills"].split(",")]
                    min_exp      = int(req_row["Min_Experience"])
                    cand_skills  = [s.strip().lower() for s in c_skills.split(",")]

                    job_desc_summary    = f"Role: {c_role}. Required Skills: {', '.join(req_skills)}. Minimum Experience: {min_exp} years."

                    # Build the best possible resume summary for the AI
                    # If PDF text is long enough, use it; otherwise build a structured summary
                    if len(resume_text.strip()) > 100:
                        resume_text_summary = resume_text[:3000]  # cap at 3000 chars
                    else:
                        # PDF extraction may have failed — build from form fields
                        resume_text_summary = (
                            f"Candidate Name: {c_name}\n"
                            f"Role Applying For: {c_role}\n"
                            f"Skills: {c_skills}\n"
                            f"Years of Experience: {c_exp} years\n"
                            f"This candidate's resume was uploaded and the above details were verified by the recruiter."
                        )

                    with st.spinner("🤖 Running AI Resume Fit Analysis — please wait..."):
                        match_result = resume_matching(
                            resume_text=resume_text_summary,
                            job_description=job_desc_summary
                        )
                        gap_result = skill_gap_analyser(
                            candidate_skills=cand_skills,
                            required_skills=req_skills,
                            experience_years=float(c_exp),
                            min_experience=min_exp
                        )

                    st.session_state["upload_match_result"] = match_result
                    st.session_state["upload_gap_result"]   = gap_result
                    st.session_state["upload_analysis_name"] = c_name
                    st.session_state["upload_analysis_role"] = c_role

        # ---- Show Analysis Results ----
        if "upload_match_result" in st.session_state and "upload_gap_result" in st.session_state:
            st.markdown("---")
            st.markdown(f"## 📊 AI Fit Analysis: **{st.session_state.get('upload_analysis_name','')}** → *{st.session_state.get('upload_analysis_role','')}*")

            # ---- Parse skill gap JSON ----
            import json, re as _re
            gap_raw    = st.session_state["upload_gap_result"]
            match_text = st.session_state["upload_match_result"]
            parsed_gap = None
            try:
                json_match = _re.search(r"(\{.*\})", gap_raw, _re.DOTALL)
                if json_match:
                    parsed_gap = json.loads(json_match.group(1))
                else:
                    cleaned = _re.sub(r"```[\w]*\n?", "", gap_raw).strip()
                    parsed_gap = json.loads(cleaned)
            except Exception:
                parsed_gap = None

            # ---- Hire Readiness Banner ----
            if parsed_gap:
                score  = parsed_gap.get("hire_readiness_score", 0)
                label  = parsed_gap.get("hire_readiness_label", "Unknown")
                exp_note = parsed_gap.get("experience_assessment", "")
                verdict  = parsed_gap.get("overall_recommendation", "")

                score_color = "#34d399" if score >= 75 else "#f59e0b" if score >= 50 else "#f87171"
                st.markdown(
                    f"""
                    <div style="background:linear-gradient(135deg,rgba(15,23,42,0.98),rgba(30,41,59,0.98));
                                border:2px solid {score_color}; border-radius:14px;
                                padding:1.2rem 1.6rem; margin:0.8rem 0;
                                display:flex; justify-content:space-between; align-items:center;">
                        <div>
                            <div style="font-size:0.8rem;color:#94a3b8;margin-bottom:0.2rem;letter-spacing:0.08em;">AI HIRE READINESS</div>
                            <div style="font-size:1.5rem;font-weight:800;color:{score_color};">{label}</div>
                            <div style="font-size:0.9rem;color:#f1f5f9;margin-top:0.3rem;font-weight:500;">{exp_note}</div>
                        </div>
                        <div style="text-align:center;">
                            <div style="font-size:2.8rem;font-weight:900;color:{score_color};">{score}</div>
                            <div style="font-size:0.72rem;color:#cbd5e1;letter-spacing:0.06em;">/ 100 READINESS SCORE</div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # Key Strengths
                strengths_raw = parsed_gap.get("key_strengths", [])
                strengths = [
                    item if isinstance(item, dict) else {"skill": str(item), "note": ""}
                    for item in strengths_raw
                ]
                if strengths:
                    st.markdown("##### ⭐ Key Strengths")
                    cols_s = st.columns(min(len(strengths), 3))
                    for i, item in enumerate(strengths):
                        with cols_s[i % 3]:
                            st.markdown(
                                f"""<div style="background:rgba(20,83,45,0.35);border:1px solid #16a34a;
                                    border-radius:10px;padding:0.85rem;margin-bottom:0.5rem;">
                                    <div style="color:#0f172a;font-weight:700;font-size:0.95rem;">✅ {item.get('skill','')}</div>
                                    <div style="color:#0f172a;font-size:0.83rem;margin-top:0.35rem;line-height:1.5;">{item.get('note','')}</div>
                                </div>""",
                                unsafe_allow_html=True
                            )

                # Growth Areas
                growth_raw = parsed_gap.get("growth_areas", [])
                growth_areas = [
                    item if isinstance(item, dict)
                    else {"skill": str(item), "severity": "Medium", "gap_note": "", "recommendation": "", "time_to_bridge": "N/A"}
                    for item in growth_raw
                ]
                if growth_areas:
                    st.markdown("##### ⚠️ Skill Gaps & Training Roadmap")
                    severity_colors = {
                        "Critical": ("#f87171", "#7f1d1d"),
                        "High":     ("#fb923c", "#7c2d12"),
                        "Medium":   ("#f59e0b", "#78350f"),
                        "Low":      ("#a78bfa", "#3b0764"),
                    }
                    for gap in growth_areas:
                        sev     = gap.get("severity", "Medium")
                        color, bg = severity_colors.get(sev, ("#94a3b8", "#1e293b"))
                        g_skill = gap.get("skill", "")
                        g_note  = gap.get("gap_note", "")
                        g_rec   = gap.get("recommendation", "")
                        g_time  = gap.get("time_to_bridge", "N/A")
                        st.markdown(
                            f"""<div style="background:rgba(15,23,42,0.85);border-left:5px solid {color};
                                border-radius:0 10px 10px 0;padding:1rem 1.3rem;margin-bottom:0.8rem;">
                                <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">
                                    <span style="color:#f8fafc;font-weight:800;">❌ {g_skill}</span>
                                    <span style="background:{bg};color:{color};border:1px solid {color};
                                          border-radius:20px;padding:0.1rem 0.6rem;font-size:0.73rem;font-weight:700;">{sev.upper()}</span>
                                </div>
                                <div style="color:#e2e8f0;font-size:0.87rem;margin-bottom:0.6rem;line-height:1.55;">{g_note}</div>
                                <div style="display:flex;gap:0.8rem;flex-wrap:wrap;">
                                    <div style="background:rgba(37,99,235,0.25);border:1px solid #3b82f6;
                                         border-radius:6px;padding:0.4rem 0.7rem;flex:1;min-width:180px;">
                                        <div style="color:#93c5fd;font-size:0.7rem;font-weight:700;margin-bottom:0.25rem;letter-spacing:0.05em;">📚 RECOMMENDED ACTION</div>
                                        <div style="color:#f1f5f9;font-size:0.82rem;line-height:1.5;">{g_rec}</div>
                                    </div>
                                    <div style="background:rgba(109,40,217,0.25);border:1px solid #a78bfa;
                                         border-radius:6px;padding:0.4rem 0.7rem;min-width:110px;text-align:center;">
                                        <div style="color:#c4b5fd;font-size:0.7rem;font-weight:700;margin-bottom:0.25rem;letter-spacing:0.05em;">⏱ TIME TO BRIDGE</div>
                                        <div style="color:#f5f3ff;font-size:0.84rem;font-weight:800;">{g_time}</div>
                                    </div>
                                </div>
                            </div>""",
                            unsafe_allow_html=True
                        )

                # Recruiter Verdict
                if verdict:
                    st.markdown(
                        f"""<div style="background:rgba(30,64,175,0.2);border:1px solid #3b82f6;
                            border-radius:12px;padding:1.1rem 1.5rem;margin-top:0.8rem;">
                            <div style="color:#93c5fd;font-weight:700;font-size:0.82rem;margin-bottom:0.4rem;letter-spacing:0.05em;">🧠 AI RECRUITER VERDICT</div>
                            <div style="color:#0f172a;font-size:0.92rem;line-height:1.7;font-weight:500;">{verdict}</div>
                        </div>""",
                        unsafe_allow_html=True
                    )
            else:
                # JSON parse failed — show raw gap output
                st.markdown("##### ⚠️ AI Skill Gap Output")
                st.info(gap_raw)

            # ---- Resume Match Summary (always show) ----
            st.markdown("---")
            st.markdown("##### 📄 AI Resume Match Summary")
            # Escape any HTML tags in the AI response to prevent layout breakage
            import html as _html
            safe_match_text = _html.escape(match_text).replace("\n", "<br>")
            st.markdown(
                f"""<div style="background:rgba(15,23,42,0.92);border:1px solid #334155;
                    border-radius:12px;padding:1.4rem 1.6rem;color:#f1f5f9;
                    font-size:0.94rem;line-height:1.8;white-space:pre-wrap;">
                    {safe_match_text}
                </div>""",
                unsafe_allow_html=True
            )



def render_resume_analysis(candidates_df, requirements_df):
    st.markdown("### 📊 AI Resume Analysis Report")
    
    cand_names = sorted(candidates_df["Name"].tolist())
    if not cand_names:
        st.warning("No candidates available for analysis.")
        return
        
    default_idx = 0
    if "last_analyzed_candidate" in st.session_state:
        lac = st.session_state["last_analyzed_candidate"]
        if lac in cand_names:
            default_idx = cand_names.index(lac)
            
    selected_name = st.selectbox("Select Candidate to Analyze", cand_names, index=default_idx)
    
    cand_data = candidates_df[candidates_df["Name"] == selected_name].iloc[0]
    role_applied = cand_data["Role Applied"]
    c_skills = [s.strip().lower() for s in cand_data["Skills"].split(",")]
    
    job_spec = requirements_df[requirements_df["Role"] == role_applied]
    if not job_spec.empty:
        req_skills = [s.strip().lower() for s in job_spec.iloc[0]["Required_Skills"].split(",")]
        min_exp = int(job_spec.iloc[0]["Min_Experience"])
    else:
        req_skills = []
        min_exp = 0
        
    matched = [s for s in req_skills if s in c_skills]
    missing = [s for s in req_skills if s not in c_skills]
    match_score = int((len(matched) / max(1, len(req_skills))) * 100)
    
    col_an_1, col_an_2 = st.columns([2, 1])
    with col_an_1:
        st.markdown(
            f"""
            <div class="explainability-card" style="margin-top: 0;">
                <h4 style="margin:0; color:#f8fafc;">Resume Score Summary: {selected_name}</h4>
                <p style="color:#94a3b8; margin:0.2rem 0 1rem 0;">Evaluated against: <strong>{role_applied}</strong></p>
                <div style="font-size:1.1rem; line-height: 1.6; color:#cbd5e1;">
                    ✔️ <strong>Matched Skills:</strong> <span style="color:#34d399;">{', '.join(matched) if matched else 'None'}</span><br>
                    ❌ <strong>Skill Gaps:</strong> <span style="color:#f87171;">{', '.join(missing) if missing else 'None'}</span><br>
                    📅 <strong>Experience Validation:</strong> Candidate has {cand_data['Experience_Years']} years (Job Requirement: {min_exp} years).
                </div>
            </div>
            """, unsafe_allow_html=True
        )
        
        st.markdown("#### 🤖 Core Resume Recommendations")
        if not missing:
            st.success("Excellent! The candidate matches all required skills for this position.")
        else:
            st.info(f"Recommended Upskilling: Suggest candidate completes certifications or projects in: **{', '.join(missing).upper()}**.")
            
        # Check cache
        if "ai_analysis_cache" not in st.session_state:
            st.session_state["ai_analysis_cache"] = {}
        cache = st.session_state["ai_analysis_cache"].get(selected_name, {})
        
        if st.button("🤖 Run AI Skill Gap Analysis", key=f"run_ai_{selected_name}", use_container_width=True, type="primary"):
            with st.spinner("Running deep AI analysis — this may take 10-20 seconds..."):
                gap_res = skill_gap_analyser(
                    candidate_skills=c_skills,
                    required_skills=req_skills,
                    experience_years=cand_data['Experience_Years'],
                    min_experience=min_exp
                )
                resume_text_summary = f"Candidate Name: {selected_name}. Role Applied: {role_applied}. Skills: {cand_data['Skills']}. Experience: {cand_data['Experience_Years']} years."
                job_desc_summary = f"Required Skills: {', '.join(req_skills)}. Minimum Experience: {min_exp} years."
                match_res = resume_matching(resume_text=resume_text_summary, job_description=job_desc_summary)
                
                cache = {
                    "skill_gap": gap_res,
                    "match_summary": match_res
                }
                st.session_state["ai_analysis_cache"][selected_name] = cache
                safe_rerun()

        # ---- Render rich AI output ----
        if cache and "skill_gap" in cache:
            raw_json = cache["skill_gap"]
            parsed = None
            try:
                import json, re
                json_match = re.search(r"(\{.*\})", raw_json, re.DOTALL)
                if json_match:
                    parsed = json.loads(json_match.group(1))
                else:
                    cleaned = re.sub(r"```[\w]*\n?", "", raw_json).strip()
                    parsed = json.loads(cleaned)
            except Exception:
                parsed = None

            if parsed:
                # ---- Hire Readiness Banner ----
                score = parsed.get("hire_readiness_score", 0)
                label = parsed.get("hire_readiness_label", "Unknown")
                exp_note = parsed.get("experience_assessment", "")
                verdict = parsed.get("overall_recommendation", "")

                score_color = (
                    "#34d399" if score >= 75 else
                    "#f59e0b" if score >= 50 else
                    "#f87171"
                )
                st.markdown(
                    f"""
                    <div style="background: linear-gradient(135deg, rgba(15,23,42,0.98), rgba(30,41,59,0.98));
                                border: 2px solid {score_color}; border-radius: 14px;
                                padding: 1.2rem 1.6rem; margin: 1rem 0; display: flex;
                                justify-content: space-between; align-items: center;">
                        <div>
                            <div style="font-size:0.85rem; color:#94a3b8; margin-bottom:0.2rem; letter-spacing:0.08em;">AI HIRE READINESS</div>
                            <div style="font-size:1.6rem; font-weight:800; color:{score_color};">{label}</div>
                            <div style="font-size:0.92rem; color:#f1f5f9; margin-top:0.3rem; font-weight:500;">{exp_note}</div>
                        </div>
                        <div style="text-align:center;">
                            <div style="font-size:3rem; font-weight:900; color:{score_color};">{score}</div>
                            <div style="font-size:0.75rem; color:#cbd5e1; letter-spacing:0.06em;">/ 100 READINESS SCORE</div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # ---- Key Strengths ----
                strengths_raw = parsed.get("key_strengths", [])
                # Normalise: model may return strings or dicts
                strengths = [
                    item if isinstance(item, dict) else {"skill": str(item), "note": ""}
                    for item in strengths_raw
                ]
                if strengths:
                    st.markdown("##### ⭐ Key Strengths")
                    cols_str = st.columns(min(len(strengths), 3))
                    for i, item in enumerate(strengths):
                        with cols_str[i % 3]:
                            skill_name = item.get("skill", str(item))
                            skill_note = item.get("note", "")
                            st.markdown(
                                f"""
                                <div style="background:rgba(20,83,45,0.35); border:1px solid #16a34a;
                                            border-radius:10px; padding:0.85rem; margin-bottom:0.5rem;">
                                    <div style="color:#0f172a; font-weight:700; font-size:0.95rem;">✅ {skill_name}</div>
                                    <div style="color:#0f172a; font-size:0.83rem; margin-top:0.35rem; line-height:1.5;">{skill_note}</div>
                                </div>
                                """,
                                unsafe_allow_html=True
                            )

                # ---- Growth Areas with Severity ----
                growth_raw = parsed.get("growth_areas", [])
                # Normalise: model may return strings or dicts
                growth_areas = [
                    item if isinstance(item, dict)
                    else {"skill": str(item), "severity": "Medium", "gap_note": "", "recommendation": "", "time_to_bridge": "N/A"}
                    for item in growth_raw
                ]
                if growth_areas:
                    st.markdown("##### ⚠️ Skill Gap Analysis & Training Roadmap")
                    severity_colors = {
                        "Critical": ("#f87171", "#7f1d1d"),
                        "High":     ("#fb923c", "#7c2d12"),
                        "Medium":   ("#f59e0b", "#78350f"),
                        "Low":      ("#a78bfa", "#3b0764"),
                    }
                    for gap in growth_areas:
                        sev = gap.get("severity", "Medium") if isinstance(gap, dict) else "Medium"
                        color, bg = severity_colors.get(sev, ("#94a3b8", "#1e293b"))
                        g_skill   = gap.get("skill", "") if isinstance(gap, dict) else str(gap)
                        g_note    = gap.get("gap_note", "") if isinstance(gap, dict) else ""
                        g_rec     = gap.get("recommendation", "") if isinstance(gap, dict) else ""
                        g_time    = gap.get("time_to_bridge", "N/A") if isinstance(gap, dict) else "N/A"
                        st.markdown(
                            f"""
                            <div style="background:rgba(15,23,42,0.85); border-left:5px solid {color};
                                        border-radius:0 10px 10px 0; padding:1rem 1.3rem; margin-bottom:0.8rem;">
                                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:0.5rem;">
                                    <span style="color:#f8fafc; font-weight:800; font-size:1.05rem;">❌ {g_skill}</span>
                                    <span style="background:{bg}; color:{color}; border:1px solid {color};
                                                 border-radius:20px; padding:0.15rem 0.7rem;
                                                 font-size:0.75rem; font-weight:700;">{sev.upper()}</span>
                                </div>
                                <div style="color:#e2e8f0; font-size:0.88rem; margin-bottom:0.7rem; line-height:1.5; font-weight:400;">{g_note}</div>
                                <div style="display:flex; gap:1rem; flex-wrap:wrap;">
                                    <div style="background:rgba(37,99,235,0.25); border:1px solid #3b82f6;
                                                border-radius:6px; padding:0.45rem 0.75rem; flex:1; min-width:200px;">
                                        <div style="color:#93c5fd; font-size:0.72rem; font-weight:700; margin-bottom:0.25rem; letter-spacing:0.05em;">📚 RECOMMENDED ACTION</div>
                                        <div style="color:#f1f5f9; font-size:0.84rem; line-height:1.5;">{g_rec}</div>
                                    </div>
                                    <div style="background:rgba(109,40,217,0.25); border:1px solid #a78bfa;
                                                border-radius:6px; padding:0.45rem 0.75rem; min-width:120px; text-align:center;">
                                        <div style="color:#c4b5fd; font-size:0.72rem; font-weight:700; margin-bottom:0.25rem; letter-spacing:0.05em;">⏱ TIME TO BRIDGE</div>
                                        <div style="color:#f5f3ff; font-size:0.88rem; font-weight:800;">{g_time}</div>
                                    </div>
                                </div>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

                # ---- Recruiter Verdict ----
                if verdict:
                    st.markdown(
                        f"""
                        <div style="background:rgba(30,64,175,0.2); border:1px solid #3b82f6;
                                    border-radius:12px; padding:1.1rem 1.5rem; margin-top:1rem;">
                            <div style="color:#93c5fd; font-weight:700; font-size:0.85rem; margin-bottom:0.5rem; letter-spacing:0.05em;">🧠 AI RECRUITER VERDICT</div>
                            <div style="color:#0f172a; font-size:0.94rem; line-height:1.7; font-weight:500;">{verdict}</div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
            else:
                # Fallback: raw text display if JSON parse fails
                st.markdown("**AI Analysis Output:**")
                st.info(raw_json)

        elif not cache:
            st.info("Click **'🤖 Run AI Skill Gap Analysis'** above to generate a full recruiter-grade AI analysis.")

        st.write("")
        with st.expander("🤖 AI Resume Match Summary", expanded=False):
            if cache and "match_summary" in cache:
                st.write(cache["match_summary"])
            else:
                st.info("Click 'Run AI Skill Gap Analysis' to generate AI Resume Match Summary.")

            
    with col_an_2:
        st.markdown("#### 🎯 Overall Match Score")
        fig = go.Figure(go.Indicator(
            mode = "gauge+number",
            value = match_score,
            domain = {'x': [0, 1], 'y': [0, 1]},
            title = {'text': "Resume Quality Index", 'font': {'color': "#f8fafc", 'size': 14}},
            gauge = {
                'axis': {'range': [0, 100], 'tickcolor': "#cbd5e1"},
                'bar': {'color': "#3b82f6"},
                'steps': [
                    {'range': [0, 50], 'color': "rgba(220, 38, 38, 0.2)"},
                    {'range': [50, 75], 'color': "rgba(245, 158, 11, 0.2)"},
                    {'range': [75, 100], 'color': "rgba(16, 185, 129, 0.2)"}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': 75
                }
            }
        ))
        fig.update_layout(
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="#f8fafc", family="Inter"),
            height=280,
            margin=dict(t=30, b=10, l=10, r=10)
        )
        st.plotly_chart(fig, use_container_width=True)

    if "resume_chat_history" not in st.session_state:
        st.session_state["resume_chat_history"] = {}
    if selected_name not in st.session_state["resume_chat_history"]:
        st.session_state["resume_chat_history"][selected_name] = []
        
    chat_history = st.session_state["resume_chat_history"][selected_name]
    
    st.markdown("---")
    st.markdown("### 💬 Ask AI About This Candidate")
    
    for msg in chat_history:
        bubble_class = "chat-bubble-user" if msg["role"] == "user" else "chat-bubble-assistant"
        st.markdown(f'<div class="{bubble_class}">{msg["content"]}</div>', unsafe_allow_html=True)
        
    chat_col1, chat_col2 = st.columns([4, 1])
    with chat_col1:
        question_input = st.text_input(
            "Ask a question about this candidate...",
            key=f"q_input_{selected_name}",
            placeholder="e.g. Does this candidate have experience leading teams?"
        )
    with chat_col2:
        ask_clicked = st.button("Ask", key=f"ask_btn_{selected_name}", use_container_width=True)
        
    if ask_clicked and question_input.strip():
        chat_history.append({"role": "user", "content": question_input.strip()})
        resume_text_summary = f"Candidate Name: {selected_name}. Role Applied: {role_applied}. Skills: {cand_data['Skills']}. Experience: {cand_data['Experience_Years']} years."
        with st.spinner("Asking AI..."):
            ans = resume_chat(candidate_context=resume_text_summary, question=question_input.strip())
        chat_history.append({"role": "assistant", "content": ans})
        safe_rerun()

def render_job_description_generator():
    st.markdown("### ✍️ AI Job Description Generator")
    st.markdown("Generate optimized, SEO-friendly job descriptions using AI parameters and immediately inject them into the screening pool.")
    
    with st.form("jd_gen_form"):
        col1, col2 = st.columns(2)
        with col1:
            title = st.text_input("Job Title", placeholder="e.g. Senior Cloud Architect")
            exp_required = st.number_input("Minimum Experience (Years)", min_value=0, max_value=15, value=4)
        with col2:
            dept = st.text_input("Department", placeholder="e.g. Engineering / Cloud Services")
            skills_req = st.text_input("Key Skills (comma-separated)", placeholder="e.g. AWS, Terraform, Docker, Python")
            
        benefits = st.text_area("Benefits & Perks", placeholder="e.g. Work from anywhere, flexible hours, health insurance")
        submit_gen = st.form_submit_button("Generate Job Description", type="primary")
        
    if submit_gen:
        if not title.strip() or not skills_req.strip():
            st.error("Please provide both Job Title and Key Skills.")
        else:
            generated_jd = f"""
# Job Opportunity: {title.strip()}

### Department: {dept.strip() if dept.strip() else 'Engineering'}
### Minimum Required Experience: {exp_required} Years

## About the Role
We are seeking a highly skilled **{title.strip()}** to join our growing organization. In this position, you will design, scale, and optimize core business systems, cooperating closely with cross-functional product stakeholders to deliver robust outcomes.

## Key Responsibilities
- Architect high-performance, fault-tolerant infrastructures.
- Leverage **{skills_req.strip()}** to develop clean, scalable modules.
- Formulate standards for system testing, deployment pipeline, and performance optimization.
- Lead and mentor junior developers inside agile sprint modules.

## Technical Qualifications
- Minimum **{exp_required} years** of industry work experience.
- Deep expertise in: **{skills_req.strip()}**.
- Strong foundation in Git code control, unit testing frameworks, and systems delivery.

## What We Offer
- **Compensation**: Highly competitive package.
- **Perks**: {benefits.strip() if benefits.strip() else 'Flexible hours, remote setup, and performance bonuses.'}
- Comprehensive career development tracks and training stipends.
            """
            
            st.markdown("---")
            st.markdown("#### Generated Job Description Preview")
            st.markdown(
                f'<div style="background-color: #1e293b; padding: 2rem; border-radius: 8px; border: 1px solid #334155; color: #f8fafc;">{generated_jd}</div>', 
                unsafe_allow_html=True
            )
            
            st.session_state["temp_generated_jd"] = {
                "Role": title.strip(),
                "Required_Skills": skills_req.strip(),
                "Min_Experience": int(exp_required)
            }
            
            # Save button outside the generation form
            st.info("To add this generated position to the jobs database, verify the preview above and click the save button below.")

    if "temp_generated_jd" in st.session_state:
        if st.button("Add Generated Job directly to Database", type="primary", use_container_width=True):
            gj = st.session_state["temp_generated_jd"]
            st.session_state["jobs_db"].append(gj)
            del st.session_state["temp_generated_jd"]
            st.success(f"Added position **{gj['Role']}** to Job Database successfully!")
            st.session_state["nav_option"] = "Job Role Management (CRUD)"
            safe_rerun()

def render_hiring_pipeline():
    st.markdown("### 🗂️ Recruitment Workflow & Hiring Pipeline")
    st.markdown("Track and update candidates status stages dynamically. Select new pipeline stages to update their progress.")
    
    stages = ["Applied", "Screening", "Interviewing", "Offered", "Hired"]
    cols = st.columns(len(stages))
    
    for idx, stage in enumerate(stages):
        with cols[idx]:
            st.markdown(
                f"""
                <div style="background-color: #0f172a; padding: 0.5rem; border-radius: 8px; text-align: center; border: 1px solid #334155; margin-bottom: 1rem;">
                    <strong style="color: #cbd5e1; font-size: 0.9rem;">{stage.upper()}</strong>
                </div>
                """, unsafe_allow_html=True
            )
            
            stage_cands = [c for c in st.session_state["candidates_db"] if c.get("Status", "Applied") == stage]
            
            if not stage_cands:
                st.markdown("<p style='text-align:center; color:#64748b; font-size:0.8rem;'>Empty</p>", unsafe_allow_html=True)
            else:
                for c in stage_cands:
                    c_name = c["Name"]
                    c_role = c["Role Applied"]
                    
                    st.markdown(
                        f"""
                        <div style="background-color: #1e293b; padding: 0.8rem; border-radius: 8px; border: 1px solid #475569; margin-bottom: 0.5rem;">
                            <strong style="color: #f8fafc; font-size: 0.85rem;">{c_name}</strong><br>
                            <span style="color: #94a3b8; font-size: 0.75rem;">{c_role}</span>
                        </div>
                        """, unsafe_allow_html=True
                    )
                    
                    new_stage = st.selectbox(
                        "Move", 
                        stages, 
                        index=stages.index(stage), 
                        key=f"move_{c_name}_{stage}",
                        label_visibility="collapsed"
                    )
                    
                    if new_stage != stage:
                        for db_c in st.session_state["candidates_db"]:
                            if db_c["Name"] == c_name:
                                db_c["Status"] = new_stage
                                break
                        
                        st.session_state["notifications"].insert(0, {
                            "Message": f"Moved {c_name} from {stage} to {new_stage}.",
                            "Time": "Just now",
                            "Read": False
                        })
                        st.success(f"Moved {c_name}!")
                        safe_rerun()

def render_interview_feedback(candidates_df):
    st.markdown("### 🖋️ Interview Feedback Portal")
    st.markdown("Submit structured scorecards and qualitative evaluations for candidates after completing interviews.")
    
    cand_names = sorted(candidates_df["Name"].tolist())
    if not cand_names:
        st.warning("No candidates available.")
        return
        
    tab_manual, tab_ai = st.tabs(["Manual Scorecard", "AI Discussion Analyzer"])
    
    with tab_manual:
        with st.form("feedback_form"):
            selected_cand = st.selectbox("Select Candidate Evaluated", cand_names)
            interviewer = st.text_input("Interviewer Name", placeholder="e.g. Rachel Adams")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                tech_rating = st.slider("Technical Capability Rating (1-5)", 1, 5, 3)
            with col2:
                comm_rating = st.slider("Communication Skills Rating (1-5)", 1, 5, 3)
            with col3:
                cult_rating = st.slider("Cultural Alignment Rating (1-5)", 1, 5, 3)
                
            feedback_comments = st.text_area("Detailed Interview Remarks")
            recommendation = st.selectbox("Final Hiring Outcome Recommendation", ["Hire", "Hold", "No Hire"])
            
            submit_feedback = st.form_submit_button("Submit Evaluation Scorecard", type="primary")
            
        if submit_feedback:
            if not interviewer.strip() or not feedback_comments.strip():
                st.error("Please enter both Interviewer Name and Comments.")
            else:
                c_row = candidates_df[candidates_df["Name"] == selected_cand].iloc[0]
                new_fb = {
                    "Name": selected_cand,
                    "Role": c_row["Role Applied"],
                    "Interviewer": interviewer.strip(),
                    "Technical": int(tech_rating),
                    "Communication": int(comm_rating),
                    "Culture": int(cult_rating),
                    "Comments": feedback_comments.strip(),
                    "Outcome": recommendation
                }
                st.session_state["interview_feedback"].append(new_fb)
                
                if recommendation == "Hire":
                    for c in st.session_state["candidates_db"]:
                        if c["Name"] == selected_cand:
                            c["Status"] = "Offered"
                            break
                    st.info(f"AI Suggestion: Candidate **{selected_cand}** status updated to 'Offered' due to 'Hire' feedback.")
                    
                st.session_state["notifications"].insert(0, {
                    "Message": f"Interview scorecard submitted by {interviewer.strip()} for {selected_cand}.",
                    "Time": "Just now",
                    "Read": False
                })
                st.success("Interview feedback logged successfully!")
                safe_rerun()
                
        st.markdown("---")
        st.markdown("#### Active Feedback Submissions")
        if len(st.session_state["interview_feedback"]) == 0:
            st.info("No interview scorecards have been recorded yet.")
        else:
            fb_df = pd.DataFrame(st.session_state["interview_feedback"])
            st.dataframe(fb_df, use_container_width=True, hide_index=True)
            
    with tab_ai:
        ai_selected_cand = st.selectbox("Select Candidate for Discussion Analysis", cand_names, key="ai_feedback_cand")
        transcript = st.text_area("Paste interview transcript or conversation notes", placeholder="e.g. Recruiter: Tell me about your background...\nCandidate: I have worked on React for 3 years...", height=200)
        
        if "ai_discussion_analysis" not in st.session_state:
            st.session_state["ai_discussion_analysis"] = {}
            
        if st.button("Analyze Discussion", key="ai_discussion_btn", use_container_width=True):
            if not transcript.strip():
                st.error("Please enter some transcript or interview notes to analyze.")
            else:
                with st.spinner("Analyzing discussion with AI..."):
                    analysis_res = job_discussion_analyzer(transcript=transcript.strip())
                st.session_state["ai_discussion_analysis"][ai_selected_cand] = analysis_res
                safe_rerun()
                
        if ai_selected_cand in st.session_state["ai_discussion_analysis"]:
            st.markdown("---")
            st.markdown("#### AI Discussion Analysis Report")
            st.markdown(st.session_state["ai_discussion_analysis"][ai_selected_cand])

def render_offer_letter_generator(candidates_df):
    st.markdown("### ✉️ AI Offer Letter & Email Generator")
    st.markdown("Generate official compensation packages, interview invitations, or candidate follow-up emails, and send them directly.")
    
    email_type = st.selectbox("Email Type", [
        "Offer Letter (AI Generated)", 
        "Offer Letter (Standard Template)", 
        "Interview Invite", 
        "Rejection", 
        "Follow-up"
    ])
    
    if "current_email_type" not in st.session_state:
        st.session_state["current_email_type"] = email_type
    if st.session_state["current_email_type"] != email_type:
        st.session_state["current_email_type"] = email_type
        if "generated_ai_email" in st.session_state:
            del st.session_state["generated_ai_email"]
            
    cands_offered = [c["Name"] for c in st.session_state["candidates_db"] if c.get("Status") in ["Offered", "Hired"]]
    
    if not cands_offered:
        st.warning("No candidates currently in 'Offered' or 'Hired' pipeline stages. You can choose from all candidates below:")
        cand_list = sorted(candidates_df["Name"].tolist())
    else:
        cand_list = cands_offered

    # Recruiter inputs candidate details
    selected_cand = st.selectbox("Select Candidate", cand_list)
    
    # Editable Candidate Email ID input (Prefills dynamically based on selection)
    default_email = f"{selected_cand.lower().replace(' ', '')}@example.com"
    candidate_email = st.text_input("Candidate Email ID", value=default_email)
    
    is_offer_letter = email_type in ["Offer Letter (AI Generated)", "Offer Letter (Standard Template)"]
    
    if is_offer_letter:
        with st.form("offer_form"):
            salary = st.text_input("Annual Gross Salary (CTC)", value="₹12,0,000")
            joining_date = st.date_input("Joining Date", datetime.date.today() + datetime.timedelta(days=14))
            reporting_mgr = st.text_input("Reporting Manager", value="Amit Sen (VP of Engineering)")
            work_mode = st.selectbox("Work Location Arrangement", ["Remote", "Hybrid (2 days Office)", "Onsite / Office"])
            generate_btn = st.form_submit_button("Generate Offer Letter", type="primary")
    else:
        with st.form("ai_email_form"):
            context = st.text_area("Optional Context / Custom Message", placeholder="e.g. Schedule for Wednesday at 10 AM, include HR round details")
            generate_btn = st.form_submit_button("Generate Email Content", type="primary")
            
    if generate_btn:
        c_row = candidates_df[candidates_df["Name"] == selected_cand].iloc[0]
        c_role = c_row["Role Applied"]
        
        if email_type == "Offer Letter (Standard Template)":
            offer_text = f"""# OFFICIAL EMPLOYMENT CONTRACT & LETTER OF OFFER

**Date**: {datetime.date.today().strftime("%B %d, %Y")}  
**Ref Number**: TN/OFF/{selected_cand.upper().replace(" ", "")[:4]}/2026  

Dear **{selected_cand}**,  

We are pleased to offer you employment with **Talent Copilot Technologies Pvt. Ltd.** in the position of **{c_role}**.  

## Core Terms of Employment:
- **Salary**: {salary} per annum inclusive of all statutory components.
- **Reporting Structure**: You will report to **{reporting_mgr}**.
- **Joining Date**: Your formal date of joining is scheduled for **{joining_date.strftime("%B %d, %Y")}**.
- **Work arrangement**: **{work_mode}**.

## Probation & Notice Period:
You will be on a probation period of three (3) months, after which your employment will be confirmed subject to performance review. Notice period is 60 days on either side.

Please signify your acceptance of this offer by signing and returning the duplicate copy of this letter.

Sincerely,  
**HR Operations Team**  
*Talent Copilot Corp*"""
            st.session_state["generated_ai_email"] = offer_text
            
            offer_record = {
                "Name": selected_cand,
                "Role": c_role,
                "Salary": salary,
                "StartDate": str(joining_date),
                "Mode": work_mode,
                "Status": "Sent"
            }
            st.session_state["offer_letters"] = [ol for ol in st.session_state["offer_letters"] if ol["Name"] != selected_cand]
            st.session_state["offer_letters"].append(offer_record)
            st.session_state["last_offered_candidate"] = selected_cand
            
        elif email_type == "Offer Letter (AI Generated)":
            with st.spinner("Generating professional AI offer letter..."):
                offer_context = (
                    f"Position: {c_role}\n"
                    f"Salary CTC: {salary}\n"
                    f"Joining Date: {joining_date.strftime('%B %d, %Y')}\n"
                    f"Reporting Manager: {reporting_mgr}\n"
                    f"Work Arrangement: {work_mode}\n"
                    f"Include details of a 3-month probation and standard corporate benefits in a highly polished business format."
                )
                email_text = email_generator(
                    candidate_name=selected_cand,
                    email_type="Employment Offer Letter",
                    context=offer_context
                )
                st.session_state["generated_ai_email"] = email_text
                
                offer_record = {
                    "Name": selected_cand,
                    "Role": c_role,
                    "Salary": salary,
                    "StartDate": str(joining_date),
                    "Mode": work_mode,
                    "Status": "Sent"
                }
                st.session_state["offer_letters"] = [ol for ol in st.session_state["offer_letters"] if ol["Name"] != selected_cand]
                st.session_state["offer_letters"].append(offer_record)
                st.session_state["last_offered_candidate"] = selected_cand
                
        else:
            with st.spinner(f"Generating {email_type}..."):
                email_text = email_generator(
                    candidate_name=selected_cand,
                    email_type=email_type,
                    context=context.strip()
                )
                st.session_state["generated_ai_email"] = email_text

    if "generated_ai_email" in st.session_state:
        st.markdown("---")
        st.markdown("#### Preview and Edit Generated Content")
        
        # Recruiter can live-edit the email body in this text area
        edited_email_body = st.text_area(
            "Review / Customize Email Body:", 
            value=st.session_state["generated_ai_email"], 
            height=350,
            key="edited_email_body_text"
        )
        
        # Save updates to session state
        st.session_state["generated_ai_email"] = edited_email_body
        
        # ---- SMTP Configuration (persisted to session_state to survive reruns) ----
        with st.expander("⚙️ SMTP Mail Server Configuration (Optional)", expanded=False):
            st.markdown("Enter your SMTP credentials to send real emails to candidates. Leave blank to use **Simulation Mode**.")
            col_smtp1, col_smtp2 = st.columns([3, 1])
            with col_smtp1:
                smtp_host = st.text_input(
                    "SMTP Host", 
                    value=st.session_state.get("smtp_host", "smtp.gmail.com"),
                    key="smtp_host_input"
                )
            with col_smtp2:
                smtp_port = st.number_input(
                    "Port", min_value=1, max_value=65535,
                    value=st.session_state.get("smtp_port", 587),
                    key="smtp_port_input"
                )
            smtp_user = st.text_input(
                "SMTP Username / Sender Email",
                value=st.session_state.get("smtp_user", ""),
                placeholder="e.g. recruiter@gmail.com",
                key="smtp_user_input"
            )
            smtp_password = st.text_input(
                "App Password",
                value=st.session_state.get("smtp_password", ""),
                type="password",
                placeholder="Gmail App Password (spaces allowed)",
                key="smtp_password_input"
            )
            if st.button("💾 Save Credentials", key="save_smtp_creds"):
                st.session_state["smtp_host"] = smtp_host
                st.session_state["smtp_port"] = int(smtp_port)
                st.session_state["smtp_user"] = smtp_user
                st.session_state["smtp_password"] = smtp_password
                st.success("SMTP credentials saved for this session!")

        # ---- Resolve final SMTP values from session_state ----
        final_smtp_host = st.session_state.get("smtp_host", "smtp.gmail.com")
        final_smtp_port = int(st.session_state.get("smtp_port", 587))
        final_smtp_user = st.session_state.get("smtp_user", "").strip()
        final_smtp_pass = st.session_state.get("smtp_password", "").strip().replace(" ", "")

        # ---- Send Email Button ----
        if st.button("📨 Send Email to Candidate", type="primary", use_container_width=True):
            if not candidate_email.strip():
                st.error("Please enter a valid Candidate Email ID.")
            else:
                subject = (
                    f"Employment Offer & Agreement – {selected_cand}"
                    if is_offer_letter
                    else f"Update on your application – {selected_cand}"
                )

                if final_smtp_user and final_smtp_pass:
                    with st.spinner("Connecting to mail server and sending…"):
                        try:
                            import smtplib
                            from email.mime.text import MIMEText
                            from email.mime.multipart import MIMEMultipart

                            msg = MIMEMultipart("alternative")
                            msg["From"]    = final_smtp_user
                            msg["To"]      = candidate_email.strip()
                            msg["Subject"] = subject
                            msg.attach(MIMEText(edited_email_body, "plain", "utf-8"))

                            if final_smtp_port == 465:
                                server = smtplib.SMTP_SSL(final_smtp_host, final_smtp_port, timeout=15)
                            else:
                                server = smtplib.SMTP(final_smtp_host, final_smtp_port, timeout=15)
                                server.ehlo()
                                server.starttls()
                                server.ehlo()

                            server.login(final_smtp_user, final_smtp_pass)
                            server.sendmail(final_smtp_user, [candidate_email.strip()], msg.as_string())
                            server.quit()

                            st.success(f"🎉 Email successfully delivered to **{candidate_email.strip()}**!")
                            st.session_state["notifications"].insert(0, {
                                "Message": f"Email sent to {candidate_email.strip()} for candidate {selected_cand}.",
                                "Time": "Just now",
                                "Read": False
                            })
                        except Exception as e:
                            st.error(f"❌ SMTP Error: {e}")
                else:
                    # Simulation Mode — no credentials provided
                    st.info(
                        f"**💡 Simulation Mode – No real email sent.**\n\n"
                        f"To send a real email:\n"
                        f"1. Expand ⚙️ SMTP Configuration above.\n"
                        f"2. Enter your Gmail + App Password and click **Save Credentials**.\n"
                        f"3. Then click **Send Email** again.\n\n"
                        f"---\n"
                        f"**Simulated delivery details:**\n"
                        f"- **To**: `{candidate_email.strip()}`\n"
                        f"- **Subject**: `{subject}`"
                    )

    if "last_offered_candidate" in st.session_state:
        loc = st.session_state["last_offered_candidate"]
        if st.button(f"Simulate Acceptance for {loc}", type="secondary", use_container_width=True):
            for c in st.session_state["candidates_db"]:
                if c["Name"] == loc:
                    c["Status"] = "Hired"
                    break
            for ol in st.session_state["offer_letters"]:
                if ol["Name"] == loc:
                    ol["Status"] = "Accepted"
                    break
            if loc not in st.session_state["onboarding_db"]:
                st.session_state["onboarding_db"][loc] = {
                    "Status": "Not Started",
                    "Tasks": {"Verify Identity": False, "IT Asset Setup": False, "HR Induction": False, "Security Compliance": False}
                }
            st.session_state["notifications"].insert(0, {
                "Message": f"Offer accepted by {loc}! Ready for onboarding.",
                "Time": "Just now",
                "Read": False
            })
            st.success(f"Candidate {loc} accepted the offer! Onboarding checklist created.")
            del st.session_state["last_offered_candidate"]
            safe_rerun()

def render_onboarding_page():
    st.markdown("### 🚀 Employee Onboarding Suite")
    st.markdown("Manage checklist milestones, IT inventory allocation, and compliance declarations for hired candidates.")
    
    hired_employees = list(st.session_state["onboarding_db"].keys())
    
    if not hired_employees:
        st.info("No candidates are currently marked as 'Hired' for onboarding configuration.")
        return
        
    selected_emp = st.selectbox("Select Onboarding Candidate", hired_employees)
    emp_details = st.session_state["onboarding_db"][selected_emp]
    tasks = emp_details["Tasks"]
    
    st.markdown(f"#### Onboarding Dashboard: **{selected_emp}**")
    
    completed_tasks = sum(1 for t in tasks.values() if t)
    total_tasks = len(tasks)
    progress_pct = int((completed_tasks / total_tasks) * 100)
    
    st.write(f"Completion Status: **{emp_details['Status']}** ({progress_pct}%)")
    st.progress(progress_pct)
    
    st.markdown("---")
    st.markdown("##### Task Checklist")
    
    t1 = st.checkbox("Identity & Credential Verification (Govt ID, Degrees)", value=tasks.get("Verify Identity", False))
    t2 = st.checkbox("IT System Provisioning (Laptop, Accounts, Email)", value=tasks.get("IT Asset Setup", False))
    t3 = st.checkbox("HR Welcome Induction & Salary Accounts Signoff", value=tasks.get("HR Induction", False))
    t4 = st.checkbox("Security, Anti-Phishing & Compliance Training Modules", value=tasks.get("Security Compliance", False))
    
    if st.button("Save Onboarding Milestones", type="primary"):
        tasks["Verify Identity"] = t1
        tasks["IT Asset Setup"] = t2
        tasks["HR Induction"] = t3
        tasks["Security Compliance"] = t4
        
        all_completed = t1 and t2 and t3 and t4
        some_completed = t1 or t2 or t3 or t4
        
        if all_completed:
            emp_details["Status"] = "Completed"
        elif some_completed:
            emp_details["Status"] = "In Progress"
        else:
            emp_details["Status"] = "Not Started"
            
        st.success(f"Successfully updated onboarding progression checklist for **{selected_emp}**!")
        safe_rerun()

def render_decision_support(candidates_df, requirements_df):
    st.markdown("### 🧠 AI Recommendation & Decision Support")
    st.markdown("AI-driven decision guidance showing matching metrics, hiring advice, risk alerts, and custom training recommendations.")
    
    roles = requirements_df["Role"].tolist()
    selected_role = st.selectbox("Select Target Job Opening", roles, key="decision_role")
    
    role_req = requirements_df[requirements_df["Role"] == selected_role].iloc[0]
    req_skills = [s.strip().lower() for s in role_req["Required_Skills"].split(",")]
    
    display_cands = candidates_df[candidates_df["Role Applied"] == selected_role]
    
    if display_cands.empty:
        st.warning("No candidates found applying for this position.")
        return
        
    recommendation_list = []
    for idx, c in display_cands.iterrows():
        c_name = c["Name"]
        c_skills = [s.strip().lower() for s in c["Skills"].split(",")]
        matched = [s for s in req_skills if s in c_skills]
        missing = [s for s in req_skills if s not in c_skills]
        
        match_pct = int((len(matched) / max(1, len(req_skills))) * 100)
        exp_years = float(c["Experience_Years"])
        
        if match_pct >= 80 and exp_years >= float(role_req["Min_Experience"]):
            grade = "A+ Highly Recommended"
            action = "Immediate Interview / Fast-Track Offer"
            risk = "Low risk. Matches role requirements exceptionally well."
        elif match_pct >= 60:
            grade = "B Recommended"
            action = "Screening Call / Evaluate Missing Skills"
            risk = f"Moderate risk. Missing skills: {', '.join(missing)}."
        else:
            grade = "C Review Required"
            action = "Hold / Keep in Pipeline"
            risk = "High risk. Significant skill or experience deficiency."
            
        recommendation_list.append({
            "Name": c_name,
            "Score": match_pct,
            "Experience": exp_years,
            "Grade": grade,
            "Action": action,
            "Risk": risk,
            "Gaps": missing
        })
        
    recommendation_list = sorted(recommendation_list, key=lambda x: x["Score"], reverse=True)
    
    st.markdown("#### Candidate Recommendation Grid")
    for r in recommendation_list:
        border_color = "#10b981" if "A+" in r["Grade"] else ("#f59e0b" if "B" in r["Grade"] else "#dc2626")
        gaps_str = ", ".join(r["Gaps"]).upper() if r["Gaps"] else "None"
        
        card_html = f"""
        <div style="background-color: #1e293b; padding: 1.2rem; border-radius: 8px; border-left: 5px solid {border_color}; border-top: 1px solid #334155; border-right: 1px solid #334155; border-bottom: 1px solid #334155; margin-bottom: 1rem;">
            <div style="display: flex; justify-content: space-between; align-items: center;">
                <strong style="font-size: 1.1rem; color: #f8fafc;">{r['Name']}</strong>
                <span style="background-color: #0f172a; border: 1px solid {border_color}; color: {border_color}; padding: 0.15rem 0.5rem; border-radius: 12px; font-size: 0.8rem; font-weight: 600;">{r['Grade']}</span>
            </div>
            <div style="color: #cbd5e1; font-size: 0.9rem; margin-top: 0.5rem;">
                Match Score: <strong>{r['Score']}%</strong> | Experience: <strong>{r['Experience']} Years</strong> (Job requires: {role_req['Min_Experience']} yrs)<br>
                💥 <strong>Gap Risk:</strong> {r['Risk']}<br>
                📋 <strong>Suggested Action:</strong> {r['Action']}
            </div>
        </div>
        """
        st.markdown(card_html, unsafe_allow_html=True)
        
        if st.button("🤖 Get AI Reasoning", key=f"ai_reason_{r['Name']}", use_container_width=True):
            with st.spinner("Generating AI recommendation reasoning..."):
                reason_text = hiring_recommendation(
                    candidate_name=r['Name'],
                    match_pct=r['Score'],
                    experience=r['Experience'],
                    grade=r['Grade'],
                    gaps=gaps_str
                )
                st.session_state[f"ai_reason_text_{r['Name']}"] = reason_text
                
        if f"ai_reason_text_{r['Name']}" in st.session_state:
            st.info(st.session_state[f"ai_reason_text_{r['Name']}"])
        
        if r["Gaps"]:
            st.markdown(f"📖 **AI Onboarding Upskill Course Suggested:** Program for onboarding **{r['Name']}**: Complete courses on *{gaps_str}*.")
        st.write("")

def render_reports_export(candidates_df, requirements_df):
    st.markdown("### 📋 Reports & Data Export Portal")
    st.markdown("Download active recruiting databases, analytics summaries, and logs to CSV files.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("##### 1. Candidates Database Report")
        st.dataframe(candidates_df[["Name", "Role Applied", "Skills", "Experience_Years", "Status"]], use_container_width=True, hide_index=True)
        
        csv_candidates = candidates_df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="Download Candidates CSV",
            data=csv_candidates,
            file_name="talent_candidates_report.csv",
            mime="text/csv",
            use_container_width=True
        )
        
    with col2:
        st.markdown("##### 2. Active Positions Report")
        st.dataframe(requirements_df, use_container_width=True, hide_index=True)
        
        csv_jobs = requirements_df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="Download Job Openings CSV",
            data=csv_jobs,
            file_name="talent_job_requirements_report.csv",
            mime="text/csv",
            use_container_width=True
        )
        
    st.markdown("---")
    st.markdown("##### 3. Interview Schedules Report")
    int_df = pd.DataFrame(st.session_state["interviews"])
    st.dataframe(int_df, use_container_width=True, hide_index=True)
    
    csv_interviews = int_df.to_csv(index=False).encode('utf-8')
    st.download_button(
        label="Download Interview Schedule CSV",
        data=csv_interviews,
        file_name="talent_interviews_report.csv",
        mime="text/csv"
    )

def render_help_support():
    st.markdown("### ❓ Help & Technical Support Center")
    st.markdown("Need assistance navigating the Talent Copilot? Expand our FAQs or submit an urgent support ticket directly to our admin panel.")
    
    st.markdown("#### Frequently Asked Questions (FAQ)")
    with st.expander("How does the Dynamic Match Score work?"):
        st.write("Our system extracts candidate skills and checks their overlap with required skills listed in the job requisition, factoring in active weights tuned in 'System Settings'.")
    with st.expander("How are bias levels calculated in the audit panel?"):
        st.write("We evaluate demographic or experience concentrations inside active shortlists. If any single group constitutes >70% of the shortlist pool, we raise a skew flag.")
    with st.expander("Can I add custom job openings?"):
        st.write("Yes! Navigate to the 'Job Role Management (CRUD)' module to create, modify, or delete roles, or generate new ones in the 'AI Job Description Generator'.")
        
    st.markdown("---")
    st.markdown("#### Submit a Support Ticket")
    with st.form("support_form", clear_on_submit=True):
        name = st.text_input("Name")
        email = st.text_input("Email")
        issue_type = st.selectbox("Issue Category", ["Feature Request", "UI Bug", "Matching Engine Query", "Database Sync", "Other"])
        desc = st.text_area("Detailed Problem Statement")
        priority = st.select_slider("Ticket Priority Level", options=["Low", "Medium", "High", "Critical"])
        
        submitted = st.form_submit_button("File Support Ticket", type="primary")
        
    if submitted:
        if not name.strip() or not desc.strip():
            st.error("Please fill in your Name and Description.")
        else:
            ticket_id = f"TC-{datetime.date.today().strftime('%Y%m')}-9827"
            st.success(f"Support Ticket successfully logged! Reference ID: **{ticket_id}**. Our engineers will follow up at **{email}** within 24 hours.")

def render_about_page():
    st.markdown("### ℹ️ About AI Recruitment & Talent Copilot")
    st.markdown("This project is built to demonstrate high-fidelity talent discovery, fair-matching screening pipelines, and onboarding analytics.")
    
    st.markdown("#### 🚀 Project Architecture")
    st.markdown(
        """
        - **Frontend & App Logic**: Streamlit Web Framework
        - **Data Processing**: Pandas Dataframes & NumPy
        - **Data Visualizations**: Plotly Graph Objects & Bar/Pie charts
        - **Database Layer**: In-Memory Streamlit Session State with CSV Sync hooks
        - **Design Elements**: Custom Tailwind/Sora fonts injections with glassmorphism styles
        """
    )
    
    st.markdown("---")
    st.markdown("#### 📜 Key Platform Highlights")
    st.markdown(
        """
        - **Dynamic Weighted Matching**: Customize the importance of Technical Skills, Experience, and Culture fit.
        - **Explainable Recommendation Logs**: Plain-language ledger audits detailing matching paths.
        - **Bias & Fairness Auditor**: Actively alerts HR to candidate pool skew values.
        """
    )

# Helper to load data safely
def load_csv_data(filepath):
    try:
        if os.path.exists(filepath):
            return pd.read_csv(filepath)
        else:
            st.error(f"Critical Error: Data file '{filepath}' is missing from the disk.")
            return None
    except Exception as e:
        st.error(f"Error reading '{filepath}': {e}")
        return None

# Load raw dataset once
candidates_df_raw = load_csv_data("data/candidates.csv")
requirements_df_raw = load_csv_data("data/job_requirements.csv")

# Initialize session state databases for CRUD operations if not already set
if "jobs_db" not in st.session_state and requirements_df_raw is not None:
    st.session_state["jobs_db"] = requirements_df_raw.to_dict(orient="records")

if "candidates_db" not in st.session_state and candidates_df_raw is not None:
    st.session_state["candidates_db"] = candidates_df_raw.to_dict(orient="records")

# Ensure shortlist is in session state
if "shortlist" not in st.session_state:
    st.session_state["shortlist"] = []

# Mock data initializations if not in state
if "interviews" not in st.session_state:
    st.session_state["interviews"] = [
        {"Name": "Arun", "Role": "Software Engineer", "Date": "2026-07-10", "Time": "10:00 AM", "Mode": "Zoom Meeting", "Status": "Scheduled"},
        {"Name": "Priya", "Role": "Data Scientist", "Date": "2026-07-11", "Time": "02:00 PM", "Mode": "Microsoft Teams", "Status": "Scheduled"},
        {"Name": "Rohan", "Role": "Product Manager", "Date": "2026-07-12", "Time": "11:30 AM", "Mode": "Zoom Meeting", "Status": "Completed"}
    ]
    
if "notifications" not in st.session_state:
    st.session_state["notifications"] = [
        {"Message": "Arun has been shortlisted for Software Engineer", "Time": "2 hours ago", "Read": False},
        {"Message": "Priya has been shortlisted for Data Scientist", "Time": "4 hours ago", "Read": False},
        {"Message": "Rohan finished the Product Manager interview", "Time": "Yesterday", "Read": True}
    ]

if "employee_db" not in st.session_state:
    st.session_state["employee_db"] = [
        {
            "Name": "Vijay Kumar", 
            "Role": "Senior Software Engineer", 
            "Department": "Engineering", 
            "Tenure": "3 Years", 
            "Performance": "Exceeds Expectations",
            "Performance_History": [
                {"date": "2022-07-10", "rating": "Meets Expectations"},
                {"date": "2023-07-12", "rating": "Exceeds Expectations"},
                {"date": "2024-07-15", "rating": "Exceeds Expectations"}
            ]
        },
        {
            "Name": "Sneha Sharma", 
            "Role": "Data Scientist", 
            "Department": "Analytics", 
            "Tenure": "1.5 Years", 
            "Performance": "Meets Expectations",
            "Performance_History": [
                {"date": "2023-03-20", "rating": "Meets Expectations"},
                {"date": "2024-03-15", "rating": "Meets Expectations"}
            ]
        },
        {
            "Name": "Ramesh Patel", 
            "Role": "Lead UX Designer", 
            "Department": "Design", 
            "Tenure": "4 Years", 
            "Performance": "Outstanding",
            "Performance_History": [
                {"date": "2021-09-05", "rating": "Meets Expectations"},
                {"date": "2022-09-10", "rating": "Exceeds Expectations"},
                {"date": "2023-09-12", "rating": "Outstanding"}
            ]
        },
        {
            "Name": "Anjali Sen", 
            "Role": "HR Specialist", 
            "Department": "Human Resources", 
            "Tenure": "2 Years", 
            "Performance": "Meets Expectations",
            "Performance_History": [
                {"date": "2022-11-01", "rating": "Needs Improvement"},
                {"date": "2023-11-05", "rating": "Meets Expectations"}
            ]
        },
        {
            "Name": "Karthik Rajan", 
            "Role": "DevOps Engineer", 
            "Department": "Engineering", 
            "Tenure": "2.5 Years", 
            "Performance": "Outstanding",
            "Performance_History": [
                {"date": "2022-08-15", "rating": "Meets Expectations"},
                {"date": "2023-08-20", "rating": "Exceeds Expectations"},
                {"date": "2024-08-22", "rating": "Outstanding"}
            ]
        },
        {
            "Name": "Pooja Hegde", 
            "Role": "Frontend Developer", 
            "Department": "Engineering", 
            "Tenure": "1 Year", 
            "Performance": "Meets Expectations",
            "Performance_History": [
                {"date": "2023-10-10", "rating": "Meets Expectations"},
                {"date": "2024-10-05", "rating": "Meets Expectations"}
            ]
        },
        {
            "Name": "Vikram Malhotra", 
            "Role": "Product Manager", 
            "Department": "Product", 
            "Tenure": "5 Years", 
            "Performance": "Outstanding",
            "Performance_History": [
                {"date": "2020-05-12", "rating": "Exceeds Expectations"},
                {"date": "2022-05-15", "rating": "Exceeds Expectations"},
                {"date": "2024-05-18", "rating": "Outstanding"}
            ]
        },
        {
            "Name": "Riya Sen", 
            "Role": "Data Analyst", 
            "Department": "Analytics", 
            "Tenure": "3 Years", 
            "Performance": "Exceeds Expectations",
            "Performance_History": [
                {"date": "2021-12-01", "rating": "Meets Expectations"},
                {"date": "2022-12-05", "rating": "Exceeds Expectations"},
                {"date": "2023-12-10", "rating": "Exceeds Expectations"}
            ]
        },
        {
            "Name": "Amit Sharma", 
            "Role": "QA Automation Engineer", 
            "Department": "Engineering", 
            "Tenure": "2 Years", 
            "Performance": "Meets Expectations",
            "Performance_History": [
                {"date": "2022-06-18", "rating": "Needs Improvement"},
                {"date": "2023-06-20", "rating": "Meets Expectations"},
                {"date": "2024-06-22", "rating": "Meets Expectations"}
            ]
        },
        {
            "Name": "Meera Nair", 
            "Role": "Talent Acquisition Specialist", 
            "Department": "Human Resources", 
            "Tenure": "4 Years", 
            "Performance": "Exceeds Expectations",
            "Performance_History": [
                {"date": "2020-11-15", "rating": "Meets Expectations"},
                {"date": "2022-11-20", "rating": "Exceeds Expectations"},
                {"date": "2024-11-22", "rating": "Exceeds Expectations"}
            ]
        }
    ]
    
if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = [
        {"role": "assistant", "content": "Hello! I am your AI Talent Assistant. Ask me questions about candidates, matching criteria, or recruitment metrics."}
    ]
    
if "settings_weights" not in st.session_state:
    st.session_state["settings_weights"] = {
        "Technical Fit": 40,
        "Experience Fit": 30,
        "Communication": 15,
        "Culture Fit": 15
    }

if "candidate_notes" not in st.session_state:
    st.session_state["candidate_notes"] = {}

if "interview_feedback" not in st.session_state:
    st.session_state["interview_feedback"] = [
        {"Name": "Arun", "Role": "Software Engineer", "Interviewer": "Sarah Jenkins", "Technical": 4, "Communication": 4, "Culture": 5, "Comments": "Strong Python knowledge and good problem-solving. Lacks some React depth but is eager to learn.", "Outcome": "Hire"},
        {"Name": "Divya", "Role": "Software Engineer", "Interviewer": "John Doe", "Technical": 3, "Communication": 5, "Culture": 4, "Comments": "Excellent communication, but needs to brush up on SQL optimization queries.", "Outcome": "Hold"}
    ]

if "onboarding_db" not in st.session_state:
    st.session_state["onboarding_db"] = {
        "Arun": {"Status": "In Progress", "Tasks": {"Verify Identity": True, "IT Asset Setup": True, "HR Induction": False, "Security Compliance": False}},
        "Divya": {"Status": "Not Started", "Tasks": {"Verify Identity": False, "IT Asset Setup": False, "HR Induction": False, "Security Compliance": False}},
        "Kiran": {"Status": "Completed", "Tasks": {"Verify Identity": True, "IT Asset Setup": True, "HR Induction": True, "Security Compliance": True}}
    }

if "offer_letters" not in st.session_state:
    st.session_state["offer_letters"] = [
        {"Name": "Kiran", "Role": "Software Engineer", "Salary": "₹12,00,000", "StartDate": "2026-08-01", "Mode": "Hybrid", "Status": "Accepted"}
    ]

# Override the candidate database loading to inject a 'Status' pipeline tracker if not already present
if "candidates_db" not in st.session_state and candidates_df_raw is not None:
    db = candidates_df_raw.to_dict(orient="records")
    for i, c in enumerate(db):
        # Evenly distribute candidates across various statuses for high-fidelity visualization
        if i % 5 == 0:
            c["Status"] = "Screening"
        elif i % 5 == 1:
            c["Status"] = "Interviewing"
        elif i % 5 == 2:
            c["Status"] = "Offered"
        elif i % 5 == 3:
            c["Status"] = "Hired"
        else:
            c["Status"] = "Applied"
    st.session_state["candidates_db"] = db
elif "candidates_db" in st.session_state:
    # Ensure all items in st.session_state["candidates_db"] have a 'Status' key
    for c in st.session_state["candidates_db"]:
        if "Status" not in c:
            c["Status"] = "Applied"

# Convert session states back to dataframes so application works seamlessly
if "jobs_db" in st.session_state and len(st.session_state["jobs_db"]) > 0:
    requirements_df = pd.DataFrame(st.session_state["jobs_db"])
else:
    requirements_df = requirements_df_raw

if "candidates_db" in st.session_state and len(st.session_state["candidates_db"]) > 0:
    candidates_df = pd.DataFrame(st.session_state["candidates_db"])
else:
    candidates_df = candidates_df_raw

# 7. Apply a custom theme via st.markdown
custom_css = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@300;400;500;600;700&family=Inter:wght@300;400;500;600;700&display=swap');

/* Main font application */
html, body, [class*="css"], .stMarkdown {
    font-family: 'Inter', 'Sora', sans-serif;
}

/* Custom header with gradient */
.header-container {
    background: linear-gradient(135deg, #0f172a 0%, #1e3a8a 100%);
    padding: 2rem;
    border-radius: 16px;
    margin-bottom: 2rem;
    border: 1px solid #2d3748;
    color: #f8fafc;
    box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.3);
}

.header-title {
    font-family: 'Sora', sans-serif;
    font-size: 2.2rem;
    font-weight: 700;
    margin: 0;
    background: linear-gradient(to right, #60a5fa, #34d399);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.header-subtitle {
    font-size: 1.1rem;
    opacity: 0.8;
    margin-top: 0.5rem;
    color: #cbd5e1;
}

/* Styled Card container for candidate details */
.candidate-card {
    background-color: #1e293b;
    border-radius: 12px;
    padding: 1.2rem 1.5rem;
    margin-bottom: 1rem;
    border: 1px solid #334155;
    transition: transform 0.2s, box-shadow 0.2s;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
}

.candidate-card:hover {
    transform: translateY(-2px);
    box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.2);
    border-color: #475569;
}

.score-badge {
    background-color: #0f172a;
    border: 1px solid #3b82f6;
    color: #60a5fa;
    padding: 0.2rem 0.6rem;
    border-radius: 20px;
    font-weight: 600;
    font-size: 0.9rem;
    display: inline-block;
}

.shortlisted-badge {
    background-color: #064e3b;
    border: 1px solid #059669;
    color: #34d399;
    padding: 0.2rem 0.6rem;
    border-radius: 20px;
    font-weight: 600;
    font-size: 0.9rem;
    display: inline-block;
}

.explainability-card {
    background-color: #1e293b;
    border-left: 5px solid #3b82f6;
    padding: 1.5rem;
    border-radius: 8px;
    margin-bottom: 1.5rem;
    border-top: 1px solid #334155;
    border-right: 1px solid #334155;
    border-bottom: 1px solid #334155;
}

.explainability-header {
    font-family: 'Sora', sans-serif;
    font-weight: 600;
    font-size: 1.3rem;
    color: #f8fafc;
    margin-bottom: 0.8rem;
    display: flex;
    justify-content: space-between;
    align-items: center;
}

.audit-metric-card {
    background-color: #1e293b;
    border: 1px solid #334155;
    padding: 1.2rem;
    border-radius: 10px;
    text-align: center;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
}

.audit-number {
    font-family: 'Sora', sans-serif;
    font-size: 2.2rem;
    font-weight: 700;
    color: #60a5fa;
}

.audit-label {
    font-size: 0.9rem;
    color: #94a3b8;
    margin-top: 0.3rem;
}

/* Chat bubble styling */
.chat-bubble-user {
    background-color: #1e3a8a;
    color: #f8fafc;
    padding: 0.8rem 1.2rem;
    border-radius: 16px 16px 4px 16px;
    margin-bottom: 0.8rem;
    max-width: 80%;
    margin-left: auto;
    text-align: right;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    border: 1px solid #3b82f6;
}

.chat-bubble-assistant {
    background-color: #1e293b;
    color: #cbd5e1;
    padding: 0.8rem 1.2rem;
    border-radius: 16px 16px 16px 4px;
    margin-bottom: 0.8rem;
    max-width: 80%;
    margin-right: auto;
    text-align: left;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    border: 1px solid #334155;
}

/* Calendar Card styling */
.calendar-card {
    background-color: #1e293b;
    padding: 1rem;
    border-radius: 10px;
    margin-bottom: 0.8rem;
    border-top: 1px solid #334155;
    border-right: 1px solid #334155;
    border-bottom: 1px solid #334155;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
}

/* Timeline/Experience Tracker */
.timeline-item {
    border-left: 2px solid #3b82f6;
    padding-left: 1.2rem;
    margin-left: 0.6rem;
    padding-bottom: 1.2rem;
    position: relative;
}

.timeline-marker {
    width: 12px;
    height: 12px;
    background-color: #60a5fa;
    border-radius: 50%;
    position: absolute;
    left: -7px;
    top: 4px;
    border: 2px solid #0f172a;
}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# Main Dashboard Layout
# Title Bar containing Logo & Main Heading
col_logo, col_title = st.columns([1, 8])

with col_logo:
    # 5. Safe Logo Loading
    if os.path.exists("assets/company_logo.png"):
        st.image("assets/company_logo.png", width=110)
    else:
        st.write("[Logo Placeholder]")

with col_title:
    st.markdown(
        textwrap.dedent("""
            <div class="header-container">
                <h1 class="header-title">AI Recruitment & Talent Copilot</h1>
                <p class="header-subtitle">Advanced HR decision-making with explainable matching and fairness auditing</p>
            </div>
        """),
        unsafe_allow_html=True
    )

# Check if data successfully loaded before continuing
if candidates_df is not None and requirements_df is not None:
    
    # Precompute all_skills so it is globally available
    all_skills = set()
    for s_list in candidates_df["Skills"].dropna():
        for skill in s_list.split(","):
            all_skills.add(skill.strip())
            
    # 1. Global Navigation in Sidebar
    st.sidebar.markdown("### Navigation Module")
    
    if "nav_option" not in st.session_state:
        st.session_state["nav_option"] = "Landing / Welcome Page"
        
    workspaces_list = [
        "Landing / Welcome Page",
        "Candidate Application Portal",
        "Resume Upload Page",
        "Resume Analysis Report",
        "Candidate Assessment Suite",
        "AI Job Description Generator",
        "Recruitment Workflow / Hiring Pipeline",
        "Interview Feedback Page",
        "AI Interview Question Generator",
        "Offer Letter Generator",
        "Employee Onboarding Page",
        "AI Recommendation / Decision Support",
        "Pipeline Dashboard",
        "Job Role Management (CRUD)",
        "Candidate Profile Browser",
        "Interview Scheduler & Calendar",
        "Notifications Center",
        "Advanced HR Analytics",
        "Employee Directory",
        "AI Chatbot Assistant",
        "Reports & Export Page",
        "Help & Support Page",
        "About Project Page",
        "System Settings"
    ]
    
    current_idx = 0
    if st.session_state["nav_option"] in workspaces_list:
        current_idx = workspaces_list.index(st.session_state["nav_option"])
        
    navigation_option = st.sidebar.selectbox(
        "Select Workspace",
        workspaces_list,
        index=current_idx,
        key="nav_selectbox"
    )
    
    if st.session_state["nav_selectbox"] != st.session_state["nav_option"]:
        st.session_state["nav_option"] = st.session_state["nav_selectbox"]
        safe_rerun()
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### Target Job Role")
    roles = requirements_df["Role"].tolist()
    
    # Get index of selected role or fallback to 0
    selected_role = st.sidebar.selectbox("Select Target Position", roles)
    
    # Extract role requirements
    role_req = requirements_df[requirements_df["Role"] == selected_role].iloc[0]
    required_skills = [s.strip() for s in role_req["Required_Skills"].split(",")]
    min_experience = int(role_req["Min_Experience"])
    
    # 2. Render Conditional Sidebar Filters ONLY for Assessment Suite
    if navigation_option == "Candidate Assessment Suite":
        st.sidebar.markdown("---")
        st.sidebar.markdown("### Filter Candidate Pool")
        
        # Get unique list of all skills across candidates for filtering
        all_skills = set()
        for s_list in candidates_df["Skills"].dropna():
            for skill in s_list.split(","):
                all_skills.add(skill.strip())
                
        filtered_skills = st.sidebar.multiselect(
            "Filter by skills (Union)", 
            options=sorted(list(all_skills)),
            help="Display candidates who possess any of these selected skills."
        )
        
        st.sidebar.markdown("---")
        # Quick reset shortlist button
        st.sidebar.markdown("### Shortlist Control")
        if st.sidebar.button("Clear Candidate Shortlist", use_container_width=True):
            st.session_state["shortlist"] = []
            st.sidebar.success("Shortlist cleared!")
            safe_rerun()
            
        st.sidebar.info(f"Currently Shortlisted: **{len(st.session_state['shortlist'])}** candidate(s)")
    else:
        filtered_skills = []

    # Calculate Dynamic Match Score & Prepare DataFrame
    candidates_df_calc = candidates_df.copy()
    
    # Perform Dynamic Overlap Keyword Match Calculation
    def calculate_score(candidate_skills_str, req_skills_list):
        if not candidate_skills_str:
            return 0
        c_skills = [s.strip().lower() for s in candidate_skills_str.split(",")]
        r_skills = [s.lower() for s in req_skills_list]
        
        matches = [s for s in c_skills if s in r_skills]
        if not r_skills:
            return 100
        return int((len(matches) / len(r_skills)) * 100)
    
    # Calculate scores with configured settings weights if available
    weights = st.session_state.get("settings_weights", {"Technical Fit": 40, "Experience Fit": 30, "Communication": 15, "Culture Fit": 15})
    
    candidates_df_calc["Match_Score"] = candidates_df_calc["Skills"].apply(
        lambda x: calculate_score(x, required_skills)
    )
    
    # Sort candidates by match score
    candidates_df_calc = candidates_df_calc.sort_values(by="Match_Score", ascending=False)
    
    # Filter candidates based on Role Applied and Skill selection
    display_candidates = candidates_df_calc[candidates_df_calc["Role Applied"] == selected_role]
    
    if filtered_skills:
        # Keep candidates who have at least one of the filtered skills
        def matches_skills(skills_str):
            c_skills = [s.strip().lower() for s in skills_str.split(",")]
            f_skills = [s.lower() for s in filtered_skills]
            return any(s in c_skills for s in f_skills)
            
        display_candidates = display_candidates[display_candidates["Skills"].apply(matches_skills)]

    # ------------------ NEW CUSTOM WORKSPACE MODULES ROUTING ------------------
    if navigation_option == "Landing / Welcome Page":
        render_landing_page(candidates_df_calc, requirements_df)
        
    elif navigation_option == "Candidate Application Portal":
        render_candidate_portal(requirements_df)
        
    elif navigation_option == "Resume Upload Page":
        render_resume_upload(requirements_df)
        
    elif navigation_option == "Resume Analysis Report":
        render_resume_analysis(candidates_df_calc, requirements_df)
        
    elif navigation_option == "AI Job Description Generator":
        render_job_description_generator()
        
    elif navigation_option == "Recruitment Workflow / Hiring Pipeline":
        render_hiring_pipeline()
        
    elif navigation_option == "Interview Feedback Page":
        render_interview_feedback(candidates_df_calc)
        
    elif navigation_option == "AI Interview Question Generator":
        st.markdown("### ❓ AI Interview Question Generator")
        st.markdown("Generate highly customized technical, behavioral, and coding questions based on the candidate's target role and required skill focus.")
        
        q_role = st.selectbox("Job Role", roles, key="q_role_select")
        q_experience = st.slider("Years of Experience", 0, 15, 3, key="q_exp_slider")
        q_skills = st.multiselect("Skills to Focus on", options=sorted(list(all_skills)), key="q_skills_select")
        q_difficulty = st.select_slider("Difficulty Level", ["Easy", "Medium", "Hard"], value="Medium", key="q_diff_slider")
        
        if st.button("Generate Questions", key="generate_questions_btn", use_container_width=True):
            with st.spinner("Generating interview questions..."):
                questions_out = interview_question_generator(
                    role=q_role,
                    experience=q_experience,
                    skills=q_skills,
                    difficulty=q_difficulty
                )
                st.session_state["generated_questions"] = questions_out
                
        if "generated_questions" in st.session_state:
            st.markdown("---")
            raw_text = st.session_state["generated_questions"]
            
            sections = {
                "Technical Questions": "",
                "HR Questions": "",
                "Coding Questions": "",
                "Scenario Questions": ""
            }
            
            current_header = None
            for line in raw_text.split("\n"):
                lower_line = line.lower()
                if "#### technical questions" in lower_line:
                    current_header = "Technical Questions"
                elif "#### hr questions" in lower_line:
                    current_header = "HR Questions"
                elif "#### coding questions" in lower_line:
                    current_header = "Coding Questions"
                elif "#### scenario questions" in lower_line:
                    current_header = "Scenario Questions"
                elif line.strip().startswith("####"):
                    current_header = None
                else:
                    if current_header:
                        sections[current_header] += line + "\n"
                        
            # If all sections are empty, we display the raw text as fallback
            if not any(sections.values()):
                st.markdown(raw_text)
            else:
                for header, content in sections.items():
                    st.markdown(f"#### {header}")
                    if content.strip():
                        st.markdown(content.strip())
                    else:
                        st.info("No questions generated for this category.")
        
    elif navigation_option == "Offer Letter Generator":
        render_offer_letter_generator(candidates_df_calc)
        
    elif navigation_option == "Employee Onboarding Page":
        render_onboarding_page()
        
    elif navigation_option == "AI Recommendation / Decision Support":
        render_decision_support(candidates_df_calc, requirements_df)
        
    elif navigation_option == "Reports & Export Page":
        render_reports_export(candidates_df_calc, requirements_df)
        
    elif navigation_option == "Help & Support Page":
        render_help_support()
        
    elif navigation_option == "About Project Page":
        render_about_page()

    # ------------------ MODULE 1: ASSESSMENT SUITE (ORIGINAL 3 TABS) ------------------
    elif navigation_option == "Candidate Assessment Suite":
        
        tab_pool, tab_explain, tab_bias = st.tabs([
            "Candidate Pool", 
            "Comparison & Explainability", 
            "Bias & Fairness Audit"
        ])
        
        # ------------------ TAB 1: CANDIDATE POOL ------------------
        with tab_pool:
            st.markdown(f"### Current Pool for **{selected_role}**")
            
            # Display Job details
            st.markdown(
                f"""
                <div style="background-color: rgba(30, 41, 59, 0.4); padding: 1rem; border-radius: 8px; border: 1px solid #334155; margin-bottom: 1.5rem;">
                    <strong>Position Requirements:</strong> Required Skills: <code>{', '.join(required_skills)}</code> | Minimum Experience: <code>{min_experience} Years</code>
                </div>
                """, 
                unsafe_allow_html=True
            )
            
            # Display tabular view
            st.markdown("#### Overview Table")
            st.dataframe(
                display_candidates[["Name", "Skills", "Experience_Years", "Match_Score"]], 
                use_container_width=True,
                hide_index=True
            )
            
            st.markdown("#### Candidate Profile Action Grid")
            if display_candidates.empty:
                st.info("No candidates match your criteria. Try adjusting the sidebar filters.")
            else:
                # Custom styled list view with shortlist button action
                for idx, row in display_candidates.iterrows():
                    name = row["Name"]
                    skills = row["Skills"]
                    experience = row["Experience_Years"]
                    score = row["Match_Score"]
                    
                    is_shortlisted = name in st.session_state["shortlist"]
                    
                    # Construct HTML dynamically without newlines/indentation to avoid markdown parser code-block trigger
                    badge_html = f'<span class="shortlisted-badge" style="margin-left: 0.5rem;">Shortlisted</span>' if is_shortlisted else ''
                    card_html = (
                        f'<div class="candidate-card">'
                        f'<div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem;">'
                        f'<span style="font-size: 1.2rem; font-weight: 600; color: #f8fafc;">{name}</span>'
                        f'<div><span class="score-badge">Match: {score}%</span>{badge_html}</div>'
                        f'</div>'
                        f'<div style="margin-bottom: 0.5rem;"><span style="color: #94a3b8; font-size: 0.9rem;">Skills:</span> '
                        f'<span style="color: #cbd5e1; font-size: 0.9rem;">{skills}</span></div>'
                        f'<div style="display: flex; justify-content: space-between; align-items: center;">'
                        f'<span style="color: #94a3b8; font-size: 0.9rem;">Experience: <strong>{experience} years</strong> (Req: {min_experience})</span>'
                        f'</div>'
                        f'</div>'
                    )
                    st.markdown(card_html, unsafe_allow_html=True)
                    
                    # Render shortlisting button under each card safely with unique key
                    col_btn_dummy, col_btn = st.columns([5, 1])
                    with col_btn:
                        if is_shortlisted:
                            if st.button("Remove", key=f"remove_{name}_{idx}", use_container_width=True, type="secondary"):
                                st.session_state["shortlist"].remove(name)
                                
                                # Add notification
                                st.session_state["notifications"].insert(0, {
                                    "Message": f"{name} removed from the {selected_role} shortlist.",
                                    "Time": "Just now",
                                    "Read": False
                                })
                                safe_rerun()
                        else:
                            if st.button("Shortlist", key=f"shortlist_{name}_{idx}", use_container_width=True, type="primary"):
                                st.session_state["shortlist"].append(name)
                                
                                # Add notification
                                st.session_state["notifications"].insert(0, {
                                    "Message": f"{name} added to the {selected_role} shortlist.",
                                    "Time": "Just now",
                                    "Read": False
                                })
                                safe_rerun()
                    st.write("") # small spacing
            
            if not display_candidates.empty:
                st.markdown("---")
                if st.button("🤖 AI-Rank This Pool", key=f"ai_rank_pool_{selected_role}", use_container_width=True):
                    lines = []
                    for _, row in display_candidates.iterrows():
                        lines.append(f"Name: {row['Name']}, Match Score: {row['Match_Score']}%, Experience: {row['Experience_Years']} years, Skills: {row['Skills']}")
                    candidates_summary = "\n".join(lines)
                    job_desc = f"Required Skills: {', '.join(required_skills)}. Minimum Experience: {min_experience} years."
                    with st.spinner("AI is ranking candidates..."):
                        rank_res = candidate_ranking(candidates_summary=candidates_summary, job_description=job_desc)
                    st.session_state[f"ai_ranking_{selected_role}"] = rank_res
                    
                if f"ai_ranking_{selected_role}" in st.session_state:
                    with st.expander("AI Ranking Rationale", expanded=True):
                        st.markdown(st.session_state[f"ai_ranking_{selected_role}"])
                    
        # ------------------ TAB 2: COMPARISON & EXPLAINABILITY ------------------
        with tab_explain:
            st.markdown("### Candidate Comparison Radar & Explainability Ledger")
            
            # Select candidates to compare from shortlist
            shortlisted_names = list(set(st.session_state["shortlist"]))
            
            st.markdown("Select candidates from your shortlist to analyze side-by-side:")
            selected_candidates = st.multiselect(
                "Select Candidates",
                options=shortlisted_names,
                default=shortlisted_names
            )
            
            if len(selected_candidates) < 2:
                st.info("Please shortlist and select at least 2 candidates in the multiselect above to generate the radar comparison and explainability ledger.")
            else:
                # 1. Plotly radar chart logic
                radar_data = []
                
                # Prepare dimensions for radar comparison
                categories = ["Technical Fit", "Communication", "Experience Fit", "Culture Fit"]
                
                for name in selected_candidates:
                    candidate_data = candidates_df_calc[candidates_df_calc["Name"] == name].iloc[0]
                    c_skills = [s.strip().lower() for s in candidate_data["Skills"].split(",")]
                    
                    # 1. Technical Fit score (keyword match ratio)
                    tech_score = calculate_score(candidate_data["Skills"], required_skills)
                    
                    # 2. Communication score (rule-based)
                    comm_keywords = ["communication", "leadership", "presentation", "writing", "english"]
                    has_comm_skill = any(k in c_skills for k in comm_keywords)
                    comm_score = 90 if has_comm_skill else (80 if candidate_data["Experience_Years"] >= 5 else 68)
                    
                    # 3. Experience Fit score
                    exp_ratio = float(candidate_data["Experience_Years"]) / max(1, min_experience)
                    exp_score = min(100, int(exp_ratio * 100))
                    
                    # 4. Culture Fit score
                    culture_keywords = ["agile", "git", "scrum", "collaboration", "teamwork", "docker"]
                    culture_matches = len([k for k in culture_keywords if k in c_skills])
                    culture_score = min(100, 70 + (culture_matches * 5))
                    
                    # Plotly trace
                    radar_data.append(go.Scatterpolar(
                        r=[tech_score, comm_score, exp_score, culture_score],
                        theta=categories,
                        fill='toself',
                        name=name
                    ))
                
                # Layout configuration
                layout = go.Layout(
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, 100]
                        )
                    ),
                    showlegend=True,
                    paper_bgcolor="rgba(0,0,0,0)",
                    plot_bgcolor="rgba(0,0,0,0)",
                    font=dict(color="#f8fafc", family="Inter"),
                    margin=dict(t=30, b=30, l=30, r=30),
                    height=450
                )
                
                fig = go.Figure(data=radar_data, layout=layout)
                
                st.plotly_chart(fig, use_container_width=True)
                
                # 2. Explainability Ledger
                st.markdown("---")
                st.markdown("### Explainability Ledger")
                st.write("Plain-language reasoning trail showing exactly how and why each candidate achieved their score:")
                
                for name in selected_candidates:
                    candidate_data = candidates_df_calc[candidates_df_calc["Name"] == name].iloc[0]
                    c_skills = [s.strip() for s in candidate_data["Skills"].split(",")]
                    
                    # Identify matched and missing required skills
                    c_skills_lower = [s.lower() for s in c_skills]
                    matched = [s for s in required_skills if s.lower() in c_skills_lower]
                    missing = [s for s in required_skills if s.lower() not in c_skills_lower]
                    
                    # Compute score factors
                    tech_score = calculate_score(candidate_data["Skills"], required_skills)
                    c_exp = int(candidate_data["Experience_Years"])
                    
                    # Construct HTML dynamically without newlines/indentation to avoid markdown parser code-block trigger
                    matched_str = ', '.join(matched) if matched else 'None'
                    missing_str = ', '.join(missing) if missing else 'None'
                    
                    exp_analysis = (
                        f"This exceeds the minimum requirement of <strong>{min_experience} years</strong> by {c_exp - min_experience} year(s)."
                        if c_exp >= min_experience else
                        f"This is short of the minimum requirement of <strong>{min_experience} years</strong> by {min_experience - c_exp} year(s)."
                    )
                    
                    derivation_logic = (
                        f"A further experience buffer of {c_exp} years validates competency level."
                        if c_exp >= min_experience else
                        "Experience shortfall is flagged, which may affect practical execution."
                    )
                    
                    explain_html = (
                        f'<div class="explainability-card">'
                        f'<div class="explainability-header">'
                        f'<span>Candidate: <strong>{name}</strong></span>'
                        f'<span class="score-badge">Final Match Score: {tech_score}%</span>'
                        f'</div>'
                        f'<ul style="color: #cbd5e1; margin-bottom: 0; padding-left: 1.2rem;">'
                        f'<li><strong>Matched Required Skills:</strong> <span style="color: #34d399; font-weight: 500;">{matched_str}</span></li>'
                        f'<li><strong>Missing Required Skills:</strong> <span style="color: #f87171; font-weight: 500;">{missing_str}</span></li>'
                        f'<li><strong>Experience Analysis:</strong> Candidate has <strong>{c_exp} years</strong> of experience. {exp_analysis}</li>'
                        f'<li><strong>Score Derivation Logic:</strong> Technical keyword matching matched <strong>{len(matched)} of {len(required_skills)}</strong> required skills ({tech_score}% weight). {derivation_logic}</li>'
                        f'</ul>'
                        f'</div>'
                    )
                    st.markdown(explain_html, unsafe_allow_html=True)
                    
        # ------------------ TAB 3: BIAS & FAIRNESS AUDIT ------------------
        with tab_bias:
            st.markdown("### Bias & Fairness Audit Panel")
            st.markdown("This panel monitors candidate shortlisting choices for experience-based concentration, helping ensure a diverse pipeline.")
            
            shortlisted_candidates = candidates_df_calc[candidates_df_calc["Name"].isin(st.session_state["shortlist"])]
            
            # Display Overview Statistics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(
                    f'<div class="audit-metric-card">'
                    f'<div class="audit-number">{len(st.session_state["shortlist"])}</div>'
                    f'<div class="audit-label">Total Shortlisted Candidates</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )
            with col2:
                st.markdown(
                    f'<div class="audit-metric-card">'
                    f'<div class="audit-number">{len(display_candidates)}</div>'
                    f'<div class="audit-label">Available Applicant Pool</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )
            with col3:
                conversion_rate = round((len(st.session_state["shortlist"]) / max(1, len(display_candidates))) * 100, 1)
                st.markdown(
                    f'<div class="audit-metric-card">'
                    f'<div class="audit-number">{conversion_rate}%</div>'
                    f'<div class="audit-label">Shortlist Conversion Ratio</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                
            st.markdown("---")
            st.markdown("#### Experience Band Distribution Analysis")
            
            if shortlisted_candidates.empty:
                st.info("Please shortlist some candidates from the **Candidate Pool** tab to view the live fairness analysis.")
            else:
                # Group into experience bands
                def categorize_exp(exp):
                    if exp <= 2:
                        return "Entry (0-2 Years)"
                    elif exp <= 5:
                        return "Mid-Level (3-5 Years)"
                    else:
                        return "Senior (6+ Years)"
                        
                shortlisted_candidates = shortlisted_candidates.copy()
                shortlisted_candidates["Experience_Band"] = shortlisted_candidates["Experience_Years"].apply(categorize_exp)
                
                # Count sizes
                counts = shortlisted_candidates["Experience_Band"].value_counts()
                total_shortlisted = len(shortlisted_candidates)
                
                # Render distributions
                bands = ["Entry (0-2 Years)", "Mid-Level (3-5 Years)", "Senior (6+ Years)"]
                
                col_bar1, col_bar2 = st.columns([1, 1])
                
                with col_bar1:
                    st.markdown("##### Current Shortlist Spread")
                    for band in bands:
                        count = counts.get(band, 0)
                        pct = (count / total_shortlisted) * 100
                        st.write(f"**{band}**: {count} candidate(s) ({pct:.1f}%)")
                        st.progress(int(pct))
                
                with col_bar2:
                    # Audit Flag Logic
                    skewed_band = None
                    skew_pct = 0.0
                    
                    for band in bands:
                        count = counts.get(band, 0)
                        pct = (count / total_shortlisted) * 100
                        if pct > 70.0:
                            skewed_band = band
                            skew_pct = pct
                            break
                    
                    if skewed_band:
                        st.warning(
                            textwrap.dedent(f"""
                                **Potential Shortlist Skew Detected!**
                                
                                The experience band **{skewed_band}** represents **{skew_pct:.1f}%** of your shortlist, which exceeds the recommended 70% threshold.
                                
                                **Ethical AI Recommendation:**
                                A highly concentrated shortlist might introduce age or tier-based bias. To ensure cognitive diversity and team growth potential, we recommend reviewing candidates in other experience brackets before submitting the final hiring recommendation.
                            """)
                        )
                    else:
                        st.success(
                            textwrap.dedent(f"""
                                **Shortlist Diversity Check Passed**
                                
                                No single experience band exceeds the 70% concentration limit. The shortlist represents a healthy, balanced spread of experience levels.
                            """)
                        )
                        
            st.markdown("")
            st.caption("**Fairness Auditor Note:** This tool evaluates simple statistical distribution skews on selected parameters. It is a rule-based heuristic check intended for decision support, not an automated legal fairness/ML audit.")

    # ------------------ MODULE 2: PIPELINE DASHBOARD ------------------
    elif navigation_option == "Pipeline Dashboard":
        st.markdown("### Pipeline Dashboard Overview")
        
        # Dashboard KPIs
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(
                f'<div class="audit-metric-card">'
                f'<div class="audit-number">{len(requirements_df)}</div>'
                f'<div class="audit-label">Active Job Openings</div>'
                f'</div>',
                unsafe_allow_html=True
            )
        with col2:
            st.markdown(
                f'<div class="audit-metric-card">'
                f'<div class="audit-number">{len(candidates_df)}</div>'
                f'<div class="audit-label">Total Registered Applicants</div>'
                f'</div>',
                unsafe_allow_html=True
            )
        with col3:
            st.markdown(
                f'<div class="audit-metric-card">'
                f'<div class="audit-number" style="color: #34d399;">{len(st.session_state["shortlist"])}</div>'
                f'<div class="audit-label">Total Shortlisted</div>'
                f'</div>',
                unsafe_allow_html=True
            )
        with col4:
            active_interviews = len([i for i in st.session_state["interviews"] if i["Status"] == "Scheduled"])
            st.markdown(
                f'<div class="audit-metric-card">'
                f'<div class="audit-number" style="color: #a78bfa;">{active_interviews}</div>'
                f'<div class="audit-label">Scheduled Interviews</div>'
                f'</div>',
                unsafe_allow_html=True
            )
            
        st.markdown("---")
        
        col_chart1, col_chart2 = st.columns(2)
        with col_chart1:
            st.markdown("#### Applicants per Job Role")
            role_counts = candidates_df["Role Applied"].value_counts().reset_index()
            role_counts.columns = ["Role", "Count"]
            
            fig = go.Figure(data=[go.Bar(
                x=role_counts["Role"],
                y=role_counts["Count"],
                marker_color="#3b82f6",
                text=role_counts["Count"],
                textposition='auto'
            )])
            fig.update_layout(
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                font=dict(color="#f8fafc", family="Inter"),
                margin=dict(t=20, b=20, l=20, r=20),
                height=300
            )
            st.plotly_chart(fig, use_container_width=True)
            
        with col_chart2:
            st.markdown("#### Dynamic Candidate Distribution (Current Role)")
            # Show shortlist vs pool comparison
            in_shortlist_cnt = len(display_candidates[display_candidates["Name"].isin(st.session_state["shortlist"])])
            remaining_cnt = len(display_candidates) - in_shortlist_cnt
            
            labels = ["Shortlisted", "Under Review"]
            values = [in_shortlist_cnt, remaining_cnt]
            
            fig = go.Figure(data=[go.Pie(
                labels=labels,
                values=values,
                hole=.4,
                marker=dict(colors=["#34d399", "#1e293b"])
            )])
            fig.update_layout(
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                font=dict(color="#f8fafc", family="Inter"),
                margin=dict(t=20, b=20, l=20, r=20),
                height=300,
                showlegend=True
            )
            st.plotly_chart(fig, use_container_width=True)
            
        st.markdown("---")
        if st.button("🤖 Generate AI Recruitment Insight", key="ai_recruitment_insight", use_container_width=True):
            rc = candidates_df["Role Applied"].value_counts()
            with st.spinner("Analyzing pipeline and generating insights..."):
                insight = recruitment_analysis(
                    total_openings=len(requirements_df),
                    total_applicants=len(candidates_df),
                    total_shortlisted=len(st.session_state["shortlist"]),
                    role_counts=rc.to_dict()
                )
            st.session_state["ai_pipeline_insight"] = insight
            
        if "ai_pipeline_insight" in st.session_state:
            st.info(st.session_state["ai_pipeline_insight"])

    # ------------------ MODULE 3: JOB MANAGEMENT (CRUD) ------------------
    elif navigation_option == "Job Role Management (CRUD)":
        st.markdown("### Job Position Database Management")
        
        # Display existing job listings
        st.markdown("#### Active Positions")
        jobs_display_df = pd.DataFrame(st.session_state["jobs_db"])
        st.dataframe(jobs_display_df, use_container_width=True, hide_index=True)
        
        st.markdown("---")
        st.markdown("#### Manage Positions")
        
        tab_add, tab_edit, tab_delete = st.tabs(["Add New Role", "Edit Role Requirements", "Delete Role"])
        
        with tab_add:
            st.markdown("##### Create a New Job Listing")
            with st.form("add_job_form", clear_on_submit=True):
                new_role_name = st.text_input("Job Role Name", placeholder="e.g. Cloud Engineer")
                new_role_skills = st.text_input("Required Skills (Comma separated)", placeholder="e.g. AWS, Python, Docker")
                new_role_exp = st.slider("Minimum Required Experience (Years)", 0, 15, 3)
                
                submitted = st.form_submit_button("Add Position to Directory")
                if submitted:
                    if not new_role_name.strip() or not new_role_skills.strip():
                        st.error("Please enter a valid Role Name and at least one Skill.")
                    else:
                        new_role = {
                            "Role": new_role_name.strip(),
                            "Required_Skills": new_role_skills.strip(),
                            "Min_Experience": int(new_role_exp)
                        }
                        st.session_state["jobs_db"].append(new_role)
                        st.success(f"Successfully added new position: **{new_role_name}**!")
                        safe_rerun()
                        
        with tab_edit:
            st.markdown("##### Modify Existing Role Criteria")
            if len(st.session_state["jobs_db"]) == 0:
                st.info("No active jobs to edit.")
            else:
                existing_roles = [j["Role"] for j in st.session_state["jobs_db"]]
                selected_edit_role = st.selectbox("Select Role to Edit", existing_roles)
                
                # Fetch details
                role_details = [j for j in st.session_state["jobs_db"] if j["Role"] == selected_edit_role][0]
                
                with st.form("edit_job_form"):
                    updated_skills = st.text_input("Required Skills", value=role_details["Required_Skills"])
                    updated_exp = st.slider("Minimum Experience (Years)", 0, 15, int(role_details["Min_Experience"]))
                    
                    edit_submitted = st.form_submit_button("Save Position Changes")
                    if edit_submitted:
                        for j in st.session_state["jobs_db"]:
                            if j["Role"] == selected_edit_role:
                                j["Required_Skills"] = updated_skills.strip()
                                j["Min_Experience"] = int(updated_exp)
                                break
                        st.success(f"Successfully updated **{selected_edit_role}** requirements!")
                        safe_rerun()
                        
        with tab_delete:
            st.markdown("##### Remove Role from Database")
            if len(st.session_state["jobs_db"]) == 0:
                st.info("No active jobs to delete.")
            else:
                selected_del_role = st.selectbox("Select Role to Remove", [j["Role"] for j in st.session_state["jobs_db"]])
                
                st.warning(f"Are you sure you want to permanently delete the **{selected_del_role}** job requirement?")
                if st.button("Confirm Delete Role", type="primary"):
                    st.session_state["jobs_db"] = [j for j in st.session_state["jobs_db"] if j["Role"] != selected_del_role]
                    st.success(f"Successfully deleted role: **{selected_del_role}**")
                    safe_rerun()

    # ------------------ MODULE 4: CANDIDATE PROFILE BROWSER ------------------
    elif navigation_option == "Candidate Profile Browser":
        st.markdown("### Candidate Profile Browser")
        
        candidate_names = sorted(candidates_df["Name"].tolist())
        selected_c_name = st.selectbox("Select Candidate Profile", candidate_names)
        
        c_row = candidates_df[candidates_df["Name"] == selected_c_name].iloc[0]
        c_skills = c_row["Skills"]
        c_exp = c_row["Experience_Years"]
        c_role = c_row["Role Applied"]
        
        # Calculate scores for candidate applied role
        applied_req = requirements_df[requirements_df["Role"] == c_role]
        if not applied_req.empty:
            applied_req = applied_req.iloc[0]
            req_skills = [s.strip() for s in applied_req["Required_Skills"].split(",")]
            m_exp = int(applied_req["Min_Experience"])
            c_score = calculate_score(c_skills, req_skills)
        else:
            c_score = 0
            m_exp = 0
            
        col_prof1, col_prof2 = st.columns([2, 1])
        
        with col_prof1:
            # Styled Header Profile Card
            badge_is_short = '<span class="shortlisted-badge" style="margin-left:0.5rem;">Shortlisted</span>' if selected_c_name in st.session_state["shortlist"] else ''
            
            card_html = (
                f'<div class="candidate-card" style="border-left: 5px solid #3b82f6;">'
                f'<div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:1rem;">'
                f'<div>'
                f'<h2 style="margin:0; color:#f8fafc;">{selected_c_name}</h2>'
                f'<p style="margin:0; color:#94a3b8; font-size:1rem;">Applied for: <strong>{c_role}</strong></p>'
                f'</div>'
                f'<div>'
                f'<span class="score-badge" style="font-size:1.1rem; padding:0.4rem 0.8rem;">Match: {c_score}%</span>'
                f'{badge_is_short}'
                f'</div>'
                f'</div>'
                f'<div style="margin-bottom:1rem;">'
                f'<span style="color:#94a3b8; font-size:0.95rem;">Contact Information:</span><br>'
                f'<span style="color:#cbd5e1;">📧 {selected_c_name.lower().replace(" ", "")}@example.com | 📱 +91 98765 43210</span>'
                f'</div>'
                f'</div>'
            )
            st.markdown(card_html, unsafe_allow_html=True)
            
            st.markdown("#### Candidate Experience Timeline")
            timeline_html = (
                f'<div class="timeline-item">'
                f'<div class="timeline-marker"></div>'
                f'<strong>Lead Specialist ({c_role})</strong> at Tech Solutions (2024 - Present)<br>'
                f'<span style="color:#94a3b8; font-size:0.9rem;">Spearheaded agile development, deployment cycles, and coordinated key cross-functional milestones.</span>'
                f'</div>'
                f'<div class="timeline-item">'
                f'<div class="timeline-marker"></div>'
                f'<strong>Associate Professional</strong> at Innovate Corp (2021 - 2024)<br>'
                f'<span style="color:#94a3b8; font-size:0.9rem;">Implemented clean coding architectures, performed unit tests, and managed git pipelines.</span>'
                f'</div>'
                f'<div class="timeline-item" style="border:none;">'
                f'<div class="timeline-marker"></div>'
                f'<strong>Junior Executive / Trainee</strong> at Systemic Systems (2020 - 2021)<br>'
                f'<span style="color:#94a3b8; font-size:0.9rem;">Assisted in writing features, debugging, and documenting codebases.</span>'
                f'</div>'
            )
            st.markdown(timeline_html, unsafe_allow_html=True)
            
        with col_prof2:
            st.markdown("#### Assessment Details")
            st.metric("Total Experience", f"{c_exp} Years", delta=f"{int(c_exp) - m_exp} Over Req" if int(c_exp) >= m_exp else f"{int(c_exp) - m_exp} Under Req")
            
            st.markdown("##### Key Skills List")
            for skill in c_skills.split(","):
                st.markdown(f"- `{skill.strip()}`")
                
            st.markdown("---")
            # Interactive assessment notes
            st.markdown("##### Internal HR Review Notes")
            saved_notes = st.session_state["candidate_notes"].get(selected_c_name, "")
            
            note_input = st.text_area("Write interviewer remarks:", value=saved_notes, key=f"notes_{selected_c_name}")
            if st.button("Save Profile Notes"):
                st.session_state["candidate_notes"][selected_c_name] = note_input
                st.success("Notes successfully updated!")

    # ------------------ MODULE 5: INTERVIEW SCHEDULER & CALENDAR ------------------
    elif navigation_option == "Interview Scheduler & Calendar":
        st.markdown("### Interview Scheduling & Calendar Center")
        
        tab_sched, tab_cal = st.tabs(["Schedule Interview", "Calendar Schedule"])
        
        with tab_sched:
            st.markdown("##### Arrange a New Interview Session")
            
            avail_candidates = sorted(candidates_df["Name"].tolist())
            selected_int_cand = st.selectbox("Select Candidate for Interview", avail_candidates)
            
            c_det = candidates_df[candidates_df["Name"] == selected_int_cand].iloc[0]
            cand_role = c_det["Role Applied"]
            
            col_date, col_time = st.columns(2)
            with col_date:
                int_date = st.date_input("Interview Date", datetime.date.today() + datetime.timedelta(days=2))
            with col_time:
                int_time = st.selectbox("Interview Time Slot", ["09:00 AM", "10:00 AM", "11:00 AM", "02:00 PM", "03:00 PM", "04:00 PM"])
                
            int_mode = st.selectbox("Interview Mode", ["Zoom Meeting", "Microsoft Teams", "In-Person Office Interview", "Phone Screening"])
            
            if st.button("Confirm Interview Schedule", type="primary"):
                new_interview = {
                    "Name": selected_int_cand,
                    "Role": cand_role,
                    "Date": str(int_date),
                    "Time": int_time,
                    "Mode": int_mode,
                    "Status": "Scheduled"
                }
                st.session_state["interviews"].append(new_interview)
                # Create notification
                st.session_state["notifications"].insert(0, {
                    "Message": f"Interview scheduled with {selected_int_cand} for {cand_role}",
                    "Time": "Just now",
                    "Read": False
                })
                st.success(f"Interview successfully scheduled for **{selected_int_cand}** on {int_date} at {int_time} via {int_mode}!")
                safe_rerun()
                
        with tab_cal:
            st.markdown("##### Upcoming Interview Schedule")
            
            if len(st.session_state["interviews"]) == 0:
                st.info("No interviews scheduled currently.")
            else:
                for idx, interview in enumerate(st.session_state["interviews"]):
                    status_color = "#34d399" if interview["Status"] == "Completed" else "#3b82f6"
                    
                    # Custom calendar cards
                    cal_card_html = (
                        f'<div class="calendar-card" style="border-left: 4px solid {status_color};">'
                        f'<div style="display:flex; justify-content:space-between; align-items:center;">'
                        f'<strong>Candidate: {interview["Name"]}</strong>'
                        f'<span style="background-color:#1e293b; color:{status_color}; border:1px solid {status_color}; padding:0.1rem 0.5rem; border-radius:10px; font-size:0.8rem; font-weight:600;">{interview["Status"]}</span>'
                        f'</div>'
                        f'<div style="color:#94a3b8; font-size:0.9rem; margin-top:0.4rem;">'
                        f'Role: <strong>{interview["Role"]}</strong><br>'
                        f'Schedule: 📅 {interview["Date"]} at {interview["Time"]} | Mode: {interview["Mode"]}'
                        f'</div>'
                        f'</div>'
                    )
                    st.markdown(cal_card_html, unsafe_allow_html=True)
                    
                    # Update status button
                    if interview["Status"] == "Scheduled":
                        col_dummy, col_done = st.columns([5, 1])
                        with col_done:
                            if st.button("Mark Completed", key=f"int_{idx}", use_container_width=True):
                                interview["Status"] = "Completed"
                                st.success("Interview status updated!")
                                safe_rerun()

    # ------------------ MODULE 6: NOTIFICATIONS CENTER ------------------
    elif navigation_option == "Notifications Center":
        st.markdown("### Notifications Center")
        
        unread_count = len([n for n in st.session_state["notifications"] if not n["Read"]])
        st.write(f"You have **{unread_count}** unread notification(s).")
        
        col_actions1, col_actions2 = st.columns([1, 4])
        with col_actions1:
            if st.button("Mark All as Read", use_container_width=True):
                for n in st.session_state["notifications"]:
                    n["Read"] = True
                safe_rerun()
        with col_actions2:
            if st.button("Clear All History", use_container_width=True):
                st.session_state["notifications"] = []
                safe_rerun()
                
        st.markdown("---")
        
        if len(st.session_state["notifications"]) == 0:
            st.info("Your notification tray is empty.")
        else:
            for n in st.session_state["notifications"]:
                read_dot = "🟢" if not n["Read"] else "⚪"
                card_border = "1px solid #3b82f6" if not n["Read"] else "1px solid #334155"
                card_bg = "rgba(59, 130, 246, 0.05)" if not n["Read"] else "#1e293b"
                
                st.markdown(
                    f"""
                    <div style="background-color:{card_bg}; padding:1rem; border-radius:8px; border:{card_border}; margin-bottom:0.8rem; display:flex; justify-content:space-between; align-items:center;">
                        <div>
                            <span style="margin-right:0.5rem;">{read_dot}</span>
                            <span style="color:#f8fafc; font-size:0.95rem;">{n['Message']}</span>
                        </div>
                        <span style="color:#94a3b8; font-size:0.85rem;">{n['Time']}</span>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

    # ------------------ MODULE 7: HR ANALYTICS & INSIGHTS ------------------
    elif navigation_option == "Advanced HR Analytics":
        st.markdown("### Advanced HR Recruitment Analytics")
        
        # 1. Row 1 charts
        col_an1, col_an2 = st.columns(2)
        with col_an1:
            st.markdown("#### Average Match Score by Applied Role")
            avg_matches = candidates_df_calc.groupby("Role Applied")["Match_Score"].mean().reset_index()
            
            fig = go.Figure(data=[go.Bar(
                x=avg_matches["Role Applied"],
                y=avg_matches["Match_Score"],
                marker_color="#10b981",
                text=[f"{val:.1f}%" for val in avg_matches["Match_Score"]],
                textposition='auto'
            )])
            fig.update_layout(
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                font=dict(color="#f8fafc", family="Inter"),
                margin=dict(t=20, b=20, l=20, r=20),
                height=300
            )
            st.plotly_chart(fig, use_container_width=True)
            
        with col_an2:
            st.markdown("#### Experience Band Spread of All Applicants")
            
            def get_band(exp):
                if exp <= 2: return "Entry (0-2 Years)"
                elif exp <= 5: return "Mid-Level (3-5 Years)"
                else: return "Senior (6+ Years)"
                
            cand_temp = candidates_df.copy()
            cand_temp["Exp_Band"] = cand_temp["Experience_Years"].apply(get_band)
            band_counts = cand_temp["Exp_Band"].value_counts().reset_index()
            band_counts.columns = ["Band", "Count"]
            
            fig = go.Figure(data=[go.Pie(
                labels=band_counts["Band"],
                values=band_counts["Count"],
                hole=.4,
                marker=dict(colors=["#3b82f6", "#a78bfa", "#f59e0b"])
            )])
            fig.update_layout(
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                font=dict(color="#f8fafc", family="Inter"),
                margin=dict(t=20, b=20, l=20, r=20),
                height=300
            )
            st.plotly_chart(fig, use_container_width=True)
            
        st.markdown("---")
        st.markdown("#### Match Score Distribution Density")
        
        # Plot score distribution histogram
        fig_hist = go.Figure(data=[go.Histogram(
            x=candidates_df_calc["Match_Score"],
            nbinsx=10,
            marker_color="#a78bfa",
            opacity=0.75
        )])
        fig_hist.update_layout(
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="#f8fafc", family="Inter"),
            xaxis=dict(title="Match Score (%)"),
            yaxis=dict(title="Number of Candidates"),
            margin=dict(t=20, b=20, l=20, r=20),
            height=280
        )
        st.plotly_chart(fig_hist, use_container_width=True)
        
        st.markdown("---")
        if st.button("🤖 Generate Organizational Talent Insight", key="ai_talent_insight", use_container_width=True):
            band_counts_dict = dict(zip(band_counts["Band"], band_counts["Count"]))
            avg_matches_dict = dict(zip(avg_matches["Role Applied"], avg_matches["Match_Score"]))
            with st.spinner("Analyzing macro-level organizational talent data..."):
                insight_res = talent_insight(
                    avg_match_by_role=avg_matches_dict,
                    experience_band_distribution=band_counts_dict
                )
            st.session_state["ai_org_talent_insight"] = insight_res
            
        if "ai_org_talent_insight" in st.session_state:
            st.markdown("#### Organizational Talent Insight")
            st.info(st.session_state["ai_org_talent_insight"])

    # ------------------ MODULE 8: EMPLOYEE DIRECTORY ------------------
    elif navigation_option == "Employee Directory":
        st.markdown("### Corporate Employee Directory")
        
        # Metrics
        col_dir1, col_dir2, col_dir3 = st.columns(3)
        with col_dir1:
            st.metric("Corporate Employees", len(st.session_state["employee_db"]))
        with col_dir2:
            depts = set([e["Department"] for e in st.session_state["employee_db"]])
            st.metric("Active Departments", len(depts))
        with col_dir3:
            st.metric("Internal Promotion Rate", "14.2%")
            
        st.markdown("---")
        
        # Filter Department
        dept_options = ["All Departments"] + sorted(list(depts))
        selected_dept = st.selectbox("Filter Employee Department", dept_options)
        
        # Filter DB
        if selected_dept == "All Departments":
            filtered_employees = st.session_state["employee_db"]
        else:
            filtered_employees = [e for e in st.session_state["employee_db"] if e["Department"] == selected_dept]
            
        st.markdown("##### Employees List")
        emp_df_display = pd.DataFrame(filtered_employees)
        if "Performance_History" in emp_df_display.columns:
            emp_df_display = emp_df_display.drop(columns=["Performance_History"])
        st.dataframe(emp_df_display, use_container_width=True, hide_index=True)
        
        st.markdown("---")
        st.markdown("#### Individual Talent Profile")
        profile_emp_names = [e["Name"] for e in st.session_state["employee_db"]]
        selected_profile_name = st.selectbox("Select Employee for Talent Profile", profile_emp_names, key="profile_emp_selectbox")
        
        emp_data = None
        for e in st.session_state["employee_db"]:
            if e["Name"] == selected_profile_name:
                emp_data = e
                break
                
        if emp_data:
            st.markdown(
                f"""
                <div style="background-color: #1e293b; padding: 1.2rem; border-radius: 8px; border: 1px solid #334155; margin-bottom: 1rem;">
                    <div style="font-size: 1.15rem; color: #f8fafc; font-weight: 700; margin-bottom: 0.5rem;">{emp_data['Name']}</div>
                    <div style="color: #cbd5e1; font-size: 0.95rem; line-height: 1.6;">
                        💼 <strong>Role:</strong> {emp_data['Role']}<br>
                        🏢 <strong>Department:</strong> {emp_data['Department']}<br>
                        📅 <strong>Tenure:</strong> {emp_data['Tenure']}<br>
                        🎯 <strong>Current Performance Appraisal Score:</strong> {emp_data['Performance']}
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )
            
            st.markdown("##### 📈 Performance Rating History")
            history_list = emp_data.get("Performance_History", [])
            sorted_history = history_list[::-1]
            history_df = pd.DataFrame(sorted_history)
            if not history_df.empty:
                history_df.columns = ["Date", "Rating"]
                st.dataframe(history_df, use_container_width=True, hide_index=True)
            else:
                st.info("No performance history recorded.")
                
            if st.button("Generate AI Talent Summary", key=f"talent_ai_{selected_profile_name}", use_container_width=True, type="primary"):
                with st.spinner("Analyzing performance history..."):
                    summary_result = talent_management_summary(
                        name=emp_data["Name"],
                        role=emp_data["Role"],
                        department=emp_data["Department"],
                        tenure=emp_data["Tenure"],
                        performance_history=history_list
                    )
                st.session_state[f"ai_talent_summary_{selected_profile_name}"] = summary_result
                
            if f"ai_talent_summary_{selected_profile_name}" in st.session_state:
                ai_text = st.session_state[f"ai_talent_summary_{selected_profile_name}"]
                if "AI service unavailable" in ai_text:
                    st.warning(ai_text)
                else:
                    st.markdown(ai_text)
        
        st.markdown("---")
        # Promotion Simulator
        st.markdown("##### Internal Transfer / Promotion Simulator")
        with st.form("promote_employee_form"):
            emp_names = [e["Name"] for e in st.session_state["employee_db"]]
            selected_emp = st.selectbox("Select Employee to Update", emp_names)
            new_title = st.text_input("New Target Role Title")
            new_performance = st.selectbox("New Performance Appraisal Score", ["Meets Expectations", "Exceeds Expectations", "Outstanding"])
            
            p_submitted = st.form_submit_button("Update Employee Profile")
            if p_submitted:
                for emp in st.session_state["employee_db"]:
                    if emp["Name"] == selected_emp:
                        if new_title.strip():
                            emp["Role"] = new_title.strip()
                        emp["Performance"] = new_performance
                        if "Performance_History" not in emp:
                            emp["Performance_History"] = []
                        emp["Performance_History"].append({
                            "date": datetime.date.today().isoformat(),
                            "rating": new_performance
                        })
                        break
                st.success(f"Successfully updated internal profile for **{selected_emp}**!")
                safe_rerun()

    # ------------------ MODULE 9: AI CHATBOT ASSISTANT ------------------
    elif navigation_option == "AI Chatbot Assistant":
        st.markdown("### 🤖 AI Talent Copilot Assistant")
        st.write("Ask your AI assistant questions regarding candidates, match scores, experience levels, or upload a document to discuss.")

        # Initialize session state variables if not present
        if "chatbot_uploaded_text" not in st.session_state:
            st.session_state["chatbot_uploaded_text"] = None
        if "chatbot_uploaded_filename" not in st.session_state:
            st.session_state["chatbot_uploaded_filename"] = None

        col_chat, col_context = st.columns([5, 3])

        with col_context:
            st.markdown("#### 📁 Chat Context & File Upload")
            
            # File Uploader
            uploaded_file = st.file_uploader(
                "Upload context file (PDF, TXT, DOCX)", 
                type=["pdf", "txt", "docx"], 
                key="chatbot_file_uploader"
            )
            
            if uploaded_file is not None:
                # Process the file if it's new
                if st.session_state["chatbot_uploaded_filename"] != uploaded_file.name:
                    file_ext = uploaded_file.name.lower().split(".")[-1]
                    raw_bytes = uploaded_file.read()
                    extracted_text = ""
                    
                    try:
                        if file_ext == "txt":
                            extracted_text = raw_bytes.decode("utf-8", errors="ignore")
                        elif file_ext == "pdf":
                            import io
                            try:
                                import pypdf
                                reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
                                extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                            except ImportError:
                                try:
                                    import PyPDF2
                                    reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
                                    extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                                except ImportError:
                                    extracted_text = f"[PDF libraries missing. Could not extract {uploaded_file.name}]"
                        elif file_ext == "docx":
                            import io
                            import docx
                            doc = docx.Document(io.BytesIO(raw_bytes))
                            extracted_text = "\n".join(p.text for p in doc.paragraphs)
                    except Exception as e:
                        extracted_text = f"[Error reading file: {e}]"
                        
                    st.session_state["chatbot_uploaded_text"] = extracted_text
                    st.session_state["chatbot_uploaded_filename"] = uploaded_file.name
                    st.success(f"Successfully processed **{uploaded_file.name}**!")
                    safe_rerun()
            
            # Show active file status
            if st.session_state["chatbot_uploaded_filename"]:
                st.markdown(
                    f"""
                    <div style="background-color:rgba(16,185,129,0.15); border:1px solid #10b981; 
                                padding:1rem; border-radius:8px; margin-bottom:1rem;">
                        <span style="color:#10b981; font-weight:700;">🟢 ACTIVE FILE CONTEXT</span><br>
                        <span style="color:#f1f5f9; font-size:0.9rem;"><strong>File:</strong> {st.session_state['chatbot_uploaded_filename']}</span><br>
                        <span style="color:#cbd5e1; font-size:0.8rem;"><strong>Size:</strong> {len(st.session_state['chatbot_uploaded_text'] or '')} characters extracted</span>
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
                
                # Preview of extracted text
                if st.session_state["chatbot_uploaded_text"]:
                    with st.expander("🔍 Preview Extracted Text", expanded=False):
                        st.text(st.session_state["chatbot_uploaded_text"][:800] + "...")
                        
                if st.button("❌ Remove File Context", use_container_width=True, type="secondary"):
                    st.session_state["chatbot_uploaded_text"] = None
                    st.session_state["chatbot_uploaded_filename"] = None
                    # Also clear uploader state in session state if key exists
                    if "chatbot_file_uploader" in st.session_state:
                        del st.session_state["chatbot_file_uploader"]
                    st.success("File context cleared.")
                    safe_rerun()
            else:
                st.info("No file uploaded. The chatbot will query the active candidate database by default. Upload a resume, JD, or notes to chat about them!")

            st.markdown("---")
            # Clear chat button
            if st.button("🗑️ Clear Chat History", use_container_width=True):
                st.session_state["chat_history"] = [
                    {"role": "assistant", "content": "Hello! I am your AI Talent Assistant. Ask me questions about candidates, matching criteria, or recruitment metrics."}
                ]
                st.success("Chat history cleared!")
                safe_rerun()

        with col_chat:
            st.markdown("#### 💬 Conversation")
            
            # Custom styled container for messaging look
            chat_container = st.container(height=450)
            with chat_container:
                for message in st.session_state["chat_history"]:
                    with st.chat_message(message["role"]):
                        st.markdown(message["content"])

            # Preset Quick Query Buttons
            st.markdown("##### ⚡ Quick Queries")
            col_q1, col_q2, col_q3 = st.columns(3)
            with col_q1:
                q1_clicked = st.button("Who is top matched Software Engineer?", use_container_width=True, key="q_btn_1")
            with col_q2:
                q2_clicked = st.button("List candidates with React skills", use_container_width=True, key="q_btn_2")
            with col_q3:
                q3_clicked = st.button("Show candidates with >= 5 yrs exp", use_container_width=True, key="q_btn_3")

            chat_input_val = st.chat_input("Ask your query here...")

            # Chat Logic triggered
            user_query = None
            if q1_clicked:
                user_query = "Who is top matched Software Engineer?"
            elif q2_clicked:
                user_query = "List candidates with React skills"
            elif q3_clicked:
                user_query = "Show candidates with >= 5 yrs exp"
            elif chat_input_val:
                user_query = chat_input_val.strip()

            if user_query:
                # Append user message
                st.session_state["chat_history"].append({"role": "user", "content": user_query})
                
                # Assemble database context
                cols_to_use = ["Name", "Role Applied", "Skills", "Experience_Years", "Match_Score"]
                candidates_context = candidates_df_calc[cols_to_use].to_string(index=False)
                
                # If there's an uploaded file, merge it into context
                if st.session_state["chatbot_uploaded_text"]:
                    context_to_send = (
                        f"Candidates Summary Table:\n{candidates_context}\n\n"
                        f"[Uploaded Document: {st.session_state['chatbot_uploaded_filename']}]\n"
                        f"{st.session_state['chatbot_uploaded_text'][:3500]}"
                    )
                else:
                    context_to_send = candidates_context

                with st.spinner("AI is thinking..."):
                    response = chatbot_query(user_query=user_query, candidates_context=context_to_send)
                        
                st.session_state["chat_history"].append({"role": "assistant", "content": response})
                safe_rerun()

    # ------------------ MODULE 10: SYSTEM SETTINGS ------------------
    elif navigation_option == "System Settings":
        st.markdown("### Talent Copilot System Settings")
        
        st.markdown("#### Dynamic Match Weight Factors")
        st.write("Tune how weights are balanced when computing the dynamic Candidate Match Score:")
        
        weights = st.session_state["settings_weights"]
        
        col_w1, col_w2 = st.columns(2)
        with col_w1:
            w_tech = st.slider("Technical Fit Weight (%)", 0, 100, int(weights["Technical Fit"]))
            w_exp = st.slider("Experience Fit Weight (%)", 0, 100, int(weights["Experience Fit"]))
        with col_w2:
            w_comm = st.slider("Communication Fit Weight (%)", 0, 100, int(weights["Communication"]))
            w_cult = st.slider("Culture Fit Weight (%)", 0, 100, int(weights["Culture Fit"]))
            
        total_weight = w_tech + w_exp + w_comm + w_cult
        st.write(f"Current weight total: **{total_weight}%**")
        
        if total_weight != 100:
            st.error("All match weight factors must sum up to exactly **100%** to save.")
        else:
            if st.button("Save Weight Configuration"):
                weights["Technical Fit"] = w_tech
                weights["Experience Fit"] = w_exp
                weights["Communication"] = w_comm
                weights["Culture Fit"] = w_cult
                st.success("Successfully updated matching algorithm weight variables!")
                safe_rerun()
                
        st.markdown("---")
        st.markdown("#### Advanced System Parameters")
        st.checkbox("Enable automated email notifications alerts to candidates", value=False)
        st.checkbox("Trigger real-time notifications on shortlist updates", value=True)
        st.slider("Fairness Audit Skew Limit Flag (%)", 50, 100, 70)
        
        st.markdown("---")
        st.markdown("#### Database Factory Reset")
        if st.button("Reset Session State Databases", type="secondary"):
            if "jobs_db" in st.session_state: del st.session_state["jobs_db"]
            if "candidates_db" in st.session_state: del st.session_state["candidates_db"]
            if "interviews" in st.session_state: del st.session_state["interviews"]
            if "notifications" in st.session_state: del st.session_state["notifications"]
            if "employee_db" in st.session_state: del st.session_state["employee_db"]
            if "chat_history" in st.session_state: del st.session_state["chat_history"]
            if "settings_weights" in st.session_state: del st.session_state["settings_weights"]
            if "candidate_notes" in st.session_state: del st.session_state["candidate_notes"]
            st.success("Database session states cleared. Reloading standard configurations.")
            safe_rerun()
        
else:
    st.error("Please configure the CSV data files under the `data/` folder to run the Recruitment Copilot.")
